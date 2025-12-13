# -*- coding: utf-8 -*-
"""
LessonAI API – FastAPI
Estrutura esperada:
...
"""
from __future__ import annotations

# =========================
# IMPORTS ORGANIZADOS
# =========================
import io
import json
import re
import hashlib
import unicodedata  # requerido por _sort_ci/_finalize_list
# --- Utils centralizados (NÃO duplicar implementações no main.py) ---
from utils_norm import (
    # Normalização / parsing
    strip_accents, norm_key_ci, norm_text,
    unique_preserve_order, finalize_list, ensure_list,
    split_multi, split_multi_qs, split_multiline,

    # Etapa/Disciplina
    normalize_etapa, normalize_disciplina_alias,

    # Mapeamentos/aliases + getters
    ALIASES, get_field_ci as _get_field_ci, keys_from_field_map,

    # Extratores canônicos (use SEMPRE estes nomes sublinhados no backend)
    row_temas as _row_temas,
    row_objetos as _row_objetos,
    row_titles as _row_titles,
    row_conteudos as _row_conteudos,
    row_habilidades as _row_habilidades,
    row_aulas as _row_aulas,

    # Filtro unificado / coletores
    match_row, filter_and_collect, filter_and_collect_habilidades,

    # Placeholders (útil para filtrar valores “— Selecione —”)
    PLACEHOLDERS,

    # Compat antiga (export/pdf)
    _row_temas_by_context as _row_temas_by_context_utils, 
)


import urllib.parse
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Iterable
from diag_backend import DEBUG_DIAGNOSTICO, log_erro_backend
from fastapi import FastAPI, Response, Body, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse
from collections import defaultdict, Counter
from xml.sax.saxutils import escape as _xml_escape
import csv
from pydantic import BaseModel, Field
import os, sys
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path: sys.path.insert(0, BASE_DIR)

# ==== SHIMS DE COMPATIBILIDADE (coloque perto dos helpers globais) ====

# FIELD_MAP obrigatório
try:
    FIELD_MAP
except NameError:
    FIELD_MAP = {}

# split_multi (evita NameError se o nome variar)
try:
    split_multi
except NameError:
    def split_multi(val):
        if val is None:
            return []
        if isinstance(val, (list, tuple, set)):
            return [str(x).strip() for x in val if str(x).strip()]
        return [x.strip() for x in str(val).split("||") if x.strip()]

# finalize_list (normaliza lista final)
try:
    finalize_list
except NameError:
    def finalize_list(vals):
        seen = set(); out = []
        for v in (vals or []):
            s = str(v).strip()
            if not s: 
                continue
            k = s.casefold()
            if k in seen: 
                continue
            seen.add(k); out.append(s)
        return out

# normalize_disciplina_alias (fallback se não houver)
try:
    normalize_disciplina_alias
except NameError:
    def normalize_disciplina_alias(etapa, disc):
        # fallback: só normaliza espaçamento/caixa
        return (disc or "").strip()

# row_* extratores: garanta nomes públicos e assinatura com FIELD_MAP
# TEMA
if "row_temas" not in globals():
    try:
        row_temas = _row_temas_by_context  # se existir o privado
    except NameError:
        def row_temas(row, field_map, etapa, disc):
            # último fallback: tenta chaves comuns de tema
            keys = [
                "Unidade Temática","Unidade temática","UNIDADE TEMÁTICA","Unidades Temáticas",
                "TEMA","Tema","tema","Práticas de linguagem","PRÁTICAS DE LINGUAGEM",
                "Práticas de Linguagem","Linguagem","linguagem","LINGUAGEM","linguagens","Linguagens","LINGUAGENS"
            ]
            v = []
            for k in keys:
                if k in row and str(row[k]).strip():
                    val = row[k]
                    v.extend(val if isinstance(val, list) else [val])
            return [str(x).strip() for x in v if str(x).strip()]

# OBJETO
if "row_objetos" not in globals():
    try:
        row_objetos = _row_objetos
    except NameError:
        def row_objetos(row, field_map, etapa, disc):
            keys = ["Objeto do Conhecimento","Objetos do Conhecimento","OBJETO","OBJETOS","Objeto","Objetos"]
            v = []
            for k in keys:
                if k in row and str(row[k]).strip():
                    val = row[k]
                    v.extend(val if isinstance(val, list) else [val])
            return [str(x).strip() for x in v if str(x).strip()]

# TÍTULOS
if "row_titles" not in globals():
    try:
        row_titles = _row_titles
    except NameError:
        def row_titles(row, field_map, etapa, disc):
            keys = ["Título da aula","TÍTULO DA AULA","Titulo da aula","Título","TÍTULO","Titulo","TÍTULO "]
            v = []
            for k in keys:
                if k in row and str(row[k]).strip():
                    val = row[k]
                    v.extend(val if isinstance(val, list) else [val])
            return [str(x).strip() for x in v if str(x).strip()]

# Fallback seguro para constantes opcionais do utils_norm
try:
    from utils_norm import OC_CANON_KEY, is_oc_key  # usados em trechos legados
except Exception:
    OC_CANON_KEY = "objeto"
    def is_oc_key(_: str) -> bool:
        return False

def _fallback__dedup_preserve_order(seq):
    seen=set(); out=[]
    for s in (seq or []):
        if not isinstance(s,str): continue
        t=str(s).strip()
        if not t: continue
        k=t.casefold()
        if k not in seen:
            seen.add(k); out.append(t)
    return out

def _fallback__sort_ci(seq):
    return sorted(seq or [], key=lambda s: unicodedata.normalize("NFKD", str(s)).casefold())


def _fallback__finalize_list(seq):
    return _fallback__sort_ci(_fallback__dedup_preserve_order(seq))

try:
    import utils as _utils
    _dedup_preserve_order = getattr(_utils, "_dedup_preserve_order", _fallback__dedup_preserve_order)
    _sort_ci              = getattr(_utils, "_sort_ci",              _fallback__sort_ci)
    _finalize_list        = getattr(_utils, "_finalize_list",        _fallback__finalize_list)
except Exception:
    _dedup_preserve_order = _fallback__dedup_preserve_order
    _sort_ci              = _fallback__sort_ci
    _finalize_list        = _fallback__finalize_list

# [PATCH-SHIM: load_rows_for] — garante a função visível no main.py
try:
    # se você já tem num utils_xxx.py, mantenha este import
    from utils_data import load_rows_for as _load_rows_for  # ajuste o nome do módulo se necessário
except Exception:
    _load_rows_for = None

# ----------------------------------------
# Loader canônico de linhas para etapa/disciplina
# ----------------------------------------
def load_rows_for(etapa: str, disciplina: str) -> list[dict]:
    """
    Carrega linhas da base correspondente à etapa/disciplina.

    Ordem de resolução:
    1) Se existir um loader externo (_load_rows_for), delega pra ele.
    2) Senão, tenta resolver via FIELD_MAP (dataset CSV).
    3) Se não houver dataset no FIELD_MAP, tenta JSON via _load_json_by.
    """

    # Importação garantida para uso local na função (se já não estiver no topo)
    from utils_norm import norm_key_ci 

    # 🔍 Normaliza chaves de lookup
    e = (etapa or "").strip()
    d = (disciplina or "").strip() # Nome da disciplina "bonito" (ex: 'Matemática')

    print(f"[load_rows_for] etapa={e!r} disciplina={d!r}")

    # *** NOVO AJUSTE: Normaliza a chave da disciplina para a busca no FIELD_MAP ***
    # norm_key_ci transforma o nome bonito (d) na chave JSON (d_key).
    if d:
        d_key = norm_key_ci(d)
    else:
        d_key = "" # Caso a disciplina seja nula
    
    print(f"[load_rows_for] chave de disciplina (d_key) = {d_key!r}")
    # **************************************************************************

    # 1) Loader externo (se existir)
    try:
        if _load_rows_for:
            # Continua usando 'e' e 'd' (o nome original) para o loader externo, se houver
            rows = _load_rows_for(e, d) or []
            print(f"[load_rows_for] usando _load_rows_for → {len(rows)} linhas")
            return rows
    except NameError:
        # _load_rows_for não existe neste contexto
        pass

    # 2) FIELD_MAP: procura dataset mapeado
    fm_e = (FIELD_MAP.get(e) or {})
    
    # 🎯 Aplica a chave normalizada (d_key) aqui!
    disc_map = (fm_e.get("disciplinas") or {}).get(d_key) or {}
    
    dataset_path = disc_map.get("dataset") or fm_e.get("dataset")

    print(f"[load_rows_for] dataset_path (FIELD_MAP) = {dataset_path!r}")

    # 2.1) Se achou dataset no FIELD_MAP → tenta CSV
    if dataset_path:
        try:
            with open(dataset_path, "r", encoding="utf-8") as f:
                rows = list(csv.DictReader(f))
            print(f"[load_rows_for] CSV carregado de {dataset_path!r} → {len(rows)} linhas")
            return rows
        except FileNotFoundError:
            print(f"[load_rows_for] ERRO: arquivo CSV não encontrado em {dataset_path!r}")
        except Exception as ex:
            print(f"[load_rows_for] ERRO ao ler CSV {dataset_path!r}: {repr(ex)}")

    # 3) Fallback: tenta JSON automático via _load_json_by
    try:
        # Continua usando 'e' e 'd' (o nome original) para o loader JSON de fallback
        js = _load_json_by(etapa=e, disciplina=d, arquivo=None)
        print(f"[load_rows_for] fallback JSON via _load_json_by → tipo={type(js)}")

        # Formatos possíveis:
        #  - {"linhas": [...]} / {"rows": [...]} / {"data": [...]}
        if isinstance(js, dict):
            rows = js.get("linhas") or js.get("rows") or js.get("data") or []
            if isinstance(rows, list):
                print(f"[load_rows_for] JSON(dict) → {len(rows)} linhas")
                return rows

        #  - lista direta
        if isinstance(js, list):
            print(f"[load_rows_for] JSON(list) → {len(js)} linhas")
            return js

    except HTTPException as http_ex:
        if http_ex.status_code != 404:
            print(f"[load_rows_for] HTTPException em _load_json_by: {http_ex.status_code} {http_ex.detail}")
            raise
        print(f"[load_rows_for] _load_json_by retornou 404 para {e}/{d}")
    except NameError:
        print("[load_rows_for] _load_json_by não está definido; ignorando fallback JSON.")
    except Exception as ex:
        print(f"[load_rows_for] ERRO inesperado em _load_json_by: {repr(ex)}")

    # Se nada deu certo → 404 claro
    print(f"[load_rows_for] NENHUM dataset encontrado para {e}/{d} → lançando 404")
    raise HTTPException(status_code=404, detail=f"dataset não mapeado para {e}/{d}")


APP_NAME    = "LessonAI API"
APP_VERSION = "2025.08"

# ---------------------------------------------------------------------
# Caminhos (AJUSTE ESTE para seu projeto local)
# ---------------------------------------------------------------------
APP_ROOT = Path(__file__).resolve().parent
DATA_DIR = APP_ROOT / "public" / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)


# Reconhece códigos BNCC/RA (uso geral)
CODIGO_RE = re.compile(r'\b[A-Z]{2,}\d{2,}[A-Z]?\d*\b', re.I)

# Lista padrão de manifests por etapa
MANIFEST_FILES = [
    "manifest_fundamental_I.json",
    "manifest_fundamental_II.json",
    "manifest_medio.json",
]

# ---------------------------------------------------------------------
# App + CORS
# ---------------------------------------------------------------------
app = FastAPI(title=APP_NAME, version=APP_VERSION)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_origin_regex=r".*",
    allow_methods=["*"],
    allow_headers=["*"],
    allow_credentials=False,
    expose_headers=["Content-Disposition"],
)

# ---------------------------------------------------------------------
# field_map.json
# ---------------------------------------------------------------------
FIELD_MAP_PATH = DATA_DIR / "field_map.json"
try:
    with FIELD_MAP_PATH.open("r", encoding="utf-8") as f:
        FIELD_MAP = json.load(f) or {}
except Exception:
    FIELD_MAP = {}

# ---------------------------------------------------------------------
# Diretório para salvar planos
# ---------------------------------------------------------------------
SAVED_DIR = (APP_ROOT / "saved_plans")
SAVED_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------
# SHIMS de compatibilidade (algumas versões antigas chamam sem FIELD_MAP)
# ---------------------------------------------------------------------
DEBUG_ERRORS = True  # deixe True até estabilizar

# === Helper para ler listas da query (singular/plural, 'A||B' ou repetidas) ===


def _qs_get_list(req: Request, base: str) -> list[str]:
    """
    Lê 'base' (ex.: 'tema') ou 'bases' (ex.: 'temas'), aceita 'A||B' ou múltiplos
    ?tema=x&tema=y. Retorna lista de strings sanitizadas.
    """
    qp = req.query_params
    raw = qp.get(base) or qp.get(f"{base}s")
    items: list[str] = []
    if raw:
        if "||" in raw:
            items = [s.strip() for s in raw.split("||") if s.strip()]
        else:
            s = raw.strip()
            if s:
                items = [s]
    multi = qp.getlist(base) or qp.getlist(f"{base}s")
    for v in multi:
        v = (v or "").strip()
        if v and v not in items:
            items.append(v)
    return items


def _raise_422(route: str, e: Exception):
    if DEBUG_ERRORS:
        raise HTTPException(422, detail=f"Falha ao processar {route}: {e.__class__.__name__}: {e}")
    raise HTTPException(422, detail=f"Falha ao processar {route}")

def _row_temas_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return row_temas(row, FIELD_MAP, etapa, disciplina)

def _row_objetos_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return _row_objetos(row, FIELD_MAP, etapa, disciplina)

def _row_conteudos_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return _row_conteudos(row, FIELD_MAP, etapa, disciplina)

def _row_titles_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return _row_titles(row, FIELD_MAP, etapa, disciplina)

def _row_aulas_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return _row_aulas(row, FIELD_MAP, etapa, disciplina)


def _row_habilidades_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    return _row_habilidades(row, FIELD_MAP, etapa, disciplina)


# ---------------------------------------------------------------------
# Utilidades de caminho/arquivos
# ---------------------------------------------------------------------
def _safe_path(rel: str) -> Path:
    """Garante que o caminho fique dentro de DATA_DIR."""
    rel = (rel or "").strip().lstrip("/\\")
    p = (DATA_DIR / rel).resolve()
    if not str(p).startswith(str(DATA_DIR)):
        raise HTTPException(status_code=400, detail="Caminho inválido.")
    return p

_INVALID = r'<>:"/\\|?*\0'
_INVALID_RE = re.compile(f"[{re.escape(_INVALID)}]")

def _strip_accents_filename(s: Optional[str],
                            keep_case: bool = False,
                            allow_dots: bool = True,
                            max_len: int = 120) -> str:
    """Normaliza string para uso seguro como nome de arquivo."""
    if not s:
        return ""
    s = unicodedata.normalize("NFKD", str(s))
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"\s+", " ", s).strip()
    s = s if keep_case else s.casefold()
    s = _INVALID_RE.sub("_", s)
    if not allow_dots:
        s = s.replace(".", "_")
    s = s.lstrip(". ").rstrip(". ")
    s = re.sub(r"_+", "_", s)
    s = re.sub(r"\s+", " ", s)
    s = s.replace(" ", "_")
    reserved = {"con","prn","aux","nul"} | {f"com{i}" for i in range(1,10)} | {f"lpt{i}" for i in range(1,10)}
    if s.lower() in reserved:
        s = f"_{s}_"
    if max_len and len(s) > max_len:
        s = s[:max_len].rstrip("._ ")
    return s or "arquivo"

def safe_filename_with_ext(name: str, default_ext: str = "") -> str:
    """Preserva extensão final; se ausente, aplica default_ext (com ponto)."""
    name = str(name or "")
    m = re.match(r"^(.*?)(\.[^./\\ ]{1,10})$", name)
    if m:
        base, ext = m.group(1), m.group(2)
        base_safe = _strip_accents_filename(base)
        ext_safe  = _strip_accents_filename(ext, allow_dots=True, keep_case=True)
        if not ext_safe.startswith("."):
            ext_safe = "." + ext_safe
        return (base_safe or "arquivo") + ext_safe
    base_safe = _strip_accents_filename(name)
    if default_ext and not default_ext.startswith("."):
        default_ext = "." + default_ext
    return (base_safe or "arquivo") + (default_ext or "")

def _strip_accents_text(s: str) -> str:
    return unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode("ascii")

# ---------------------------------------------------------------------
# Loader robusto do JSON por etapa/disciplina/arquivo
# ---------------------------------------------------------------------
def _load_json_by(etapa: str = None, disciplina: str = None, arquivo: str = None):
    """
    Abre o JSON da disciplina:
    - Se 'arquivo' vier (ex.: 'medio/Matemática.json'), usa-o relativo a DATA_DIR.
    - Caso contrário, tenta localizar por variações de nome com/sem underline/acentos.
    """
    base = DATA_DIR
    path: Optional[Path] = None

    if arquivo:
        p = Path(arquivo)
        path = p if p.is_absolute() else (base / arquivo)
        if not path.exists() and etapa and p.name:
            path = base / etapa / p.name
    else:
        # Tenta localizar por variações do nome da disciplina
        if etapa and disciplina:
            etapa_dir = base / etapa
            variantes = {
                disciplina,
                disciplina.replace(" ", "_"),
                disciplina.replace("_", " "),
            }
            sem_acento = _strip_accents_filename(disciplina)
            variantes |= {
                sem_acento,
                sem_acento.replace(" ", "_"),
                sem_acento.replace("_", " "),
            }
            for cand in variantes:
                p = etapa_dir / f"{cand}.json"
                if p.exists():
                    path = p
                    break

    if not path or not path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"Arquivo da disciplina '{disciplina}' não encontrado (etapa={etapa})."
        )

    try:
        with path.open("r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"JSON inválido: {e}")
def _norm_ci(s: str) -> str:
    s = (s or "").replace("_", " ")
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    return re.sub(r"\s+", " ", s).strip().casefold()


# ---------------------------------------------------------------------
# Helpers diversos (sem duplicatas)
# ---------------------------------------------------------------------

def _same_etapa(row_etapa: str, req_etapa: str | None) -> bool:
    if not row_etapa:
        return True  # linha sem etapa explícita passa
    return normalize_etapa(row_etapa) == (req_etapa or normalize_etapa(row_etapa))


def _unique_preserve_order(items: Iterable[Any]) -> List[Any]:
    seen, out = set(), []
    for x in items or []:
        if x not in seen:
            seen.add(x)
            out.append(x)
    return out

def _norm_key(s: str) -> str:
    """Chave normalizada usada em de-duplicações locais (sem conflitar com utils_norm)."""
    s = unicodedata.normalize("NFD", s or "")
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    s = re.sub(r"[_\-\./]+", " ", s)
    s = re.sub(r"\s*&\s*", " e ", s)
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s

IGNORE_PAT = re.compile(r"^(descritores|descritoreslp|descritoresmt|metodologias|metodologias_conver)$", re.I)

def _dedupe_disciplinas(nomes: Iterable[str]) -> List[str]:
    """Dedupe por chave normalizada e prefere rótulo sem underline."""
    seen = {}
    for nome in nomes or []:
        key = _norm_key(nome)
        if IGNORE_PAT.match(key):
            continue
        rotulo = (nome or "").replace("_", " ")
        cur = seen.get(key)
        if cur is None or ("_" in cur and "_" not in nome):
            seen[key] = rotulo
    return sorted(seen.values(), key=lambda x: unicodedata.normalize("NFKD", x).casefold())

def _as_list(v: Any) -> List[str]:
    if v is None:
        return []
    if isinstance(v, (list, tuple, set)):
        return [str(x) for x in v if str(x).strip()]
    s = str(v)
    if "||" in s:
        return [p for p in (t.strip() for t in s.split("||")) if p]
    return [s] if s.strip() else []

def _as_bool(v: Any) -> bool:
    return str(v).strip().lower() in {"1","true","t","yes","y","on","sim"}

def _match_any(texto: Any, agulhas: Iterable[Any], *, contains: bool=False) -> bool:
    T = _norm_key(texto)
    needles = [_norm_key(n) for n in (agulhas or []) if str(n).strip()]
    if not needles:
        return True
    for N in needles:
        if contains:
            if N and N in T:
                return True
        else:
            if T == N:
                return True
    return False

def _iter_field_values(row: dict, keys: List[str]) -> List[str]:
    """Extrai valores de possíveis chaves; sempre retorna lista de strings."""
    for k in keys:
        if k in row:
            v = row[k]
            if isinstance(v, (list, tuple, set)):
                return [str(x) for x in v if str(x).strip()]
            return [str(v)] if str(v).strip() else []
    return []

def _parse_multi_param(val: str | None) -> set[str]:
    if not val:
        return set()
    decoded = urllib.parse.unquote_plus(val)
    parts = [p.strip() for p in decoded.split("||")]
    return {_norm_key(p) for p in parts if p}

# ---------------------------------------------------------------------
# Aulas – chaves corretas
# ---------------------------------------------------------------------
AULA_KEYS = [
    "aula", "Aula", "AULA", "Nº Aula", "Número da aula", "Numero da aula",
    "num_aula", "n_aula", "lesson", "Lesson"
]

def _row_aula_crua(row: dict) -> Optional[str]:
    """Extrai identificação/número da aula a partir de múltiplas chaves possíveis."""
    val = _get_field_ci(row, AULA_KEYS)
    return str(val).strip() if val else None


def _norm_ci(s: str) -> str:
    return _norm_key_ci(s or "")  # supondo que _norm_key_ci exista no seu utils_norm

def _val_matches(val: str, keys: list[str], contains: bool) -> bool:
    if not keys:
        return True
    v = _norm_ci(val)
    ks = [_norm_ci(k) for k in keys]
    if contains:
        return any(k in v for k in ks)
    return v in ks

def _row_accept(row: dict, *, tema, objeto, titulo, conteudo, aula, contains: bool,
                etapa: str | None = None, disciplina: str | None = None) -> bool:
    temas_sel  = set(tema or [])
    objs_sel   = set(objeto or [])
    tits_sel   = set(titulo or [])
    cont_sel   = set(conteudo or [])
    aulas_sel  = set(aula or [])
    return match_row(
        row, FIELD_MAP,
        temas_sel, objs_sel, tits_sel, cont_sel, aulas_sel,
        contains, etapa=etapa, disciplina=disciplina
    )

# ---------------------------------------------------------------------
# Modelos e normalização do payload de export
# ---------------------------------------------------------------------
class LinhaPlano(BaseModel):
    habilidade: str = Field('', description="Texto da habilidade/BNCC")
    codigo: str = Field('', description="Código BNCC (ex.: EF08MA04)")
    objetos_do_conhecimento: List[str] = Field(default_factory=list, description="Objetos do conhecimento")
    titulos: List[str] = Field(default_factory=list)
    objetivo_aprendizagem: str = Field('', description="Objetivo(s) de aprendizagem (gerado)")

class ExportPayload(BaseModel):
    identificacao: Optional[str] = None
    etapa: Optional[str] = None
    disciplina: Optional[str] = None

    # filtros principais
    temas: List[str] = Field(default_factory=list)
    conteudos: List[str] = Field(default_factory=list)
    titulos_da_aula: List[str] = Field(default_factory=list)
    habilidades: List[str] = Field(default_factory=list)

    # adicionais usados no seu fluxo
    aulas: List[str] = Field(default_factory=list)
    conhecimentos_previos: List[str] = Field(default_factory=list)
    metodologia_estrategias: List[str] = Field(default_factory=list)

    # linhas estruturadas (quando houver)
    linhas: List[LinhaPlano] = Field(default_factory=list)

    # payload “livre” que vem do front
    bruto: Optional[Dict[str, Any]] = None

def _normalize_lines(payload: ExportPayload) -> List[LinhaPlano]:
    """Dedup e saneamento das linhas (garante ao menos 1)."""
    if payload.linhas:
        norm: List[LinhaPlano] = []
        for ln in payload.linhas:
            objs = [str(s).strip() for s in (ln.objetos_do_conhecimento or []) if str(s).strip()]
            tits = [str(s).strip() for s in (ln.titulos or []) if str(s).strip()]
            dedup_objs = list(dict.fromkeys(objs))
            dedup_tits = list(dict.fromkeys(tits))
            norm.append(LinhaPlano(
                habilidade=(ln.habilidade or '').strip(),
                codigo=(ln.codigo or '').strip(),
                objetos_do_conhecimento=dedup_objs,
                titulos=dedup_tits,
                objetivo_aprendizagem=(ln.objetivo_aprendizagem or '').strip()
            ))
        return norm
    return [LinhaPlano()]  # fallback mínimo

def _to_legacy_dict(p: ExportPayload) -> Dict[str, Any]:
    """
    Converte ExportPayload para o dicionário legado usado pelos exports.
    Só preenche se houver valor.
    """
    out: Dict[str, Any] = {}

    if p.identificacao: out["identificacao"] = p.identificacao
    if p.etapa: out["etapa"] = p.etapa
    if p.disciplina:
        out["disciplina"] = p.disciplina
        out["componente_curricular"] = p.disciplina  # espelho

    if p.temas:
        out["tema"] = p.temas
        out["temas"] = p.temas
    if p.conteudos:
        out["conteudos"] = p.conteudos
    if p.titulos_da_aula:
        out["titulos_da_aula"] = p.titulos_da_aula
    if p.conhecimentos_previos:
        out["conhecimentos_previos"] = p.conhecimentos_previos
    if p.metodologia_estrategias:
        out["metodologia_estrategias"] = p.metodologia_estrategias

    if p.linhas:
        out["linhas"] = [
            {
                "habilidade": lp.habilidade,
                "codigo": lp.codigo,
                "objetos_do_conhecimento": getattr(lp, "objetos_do_conhecimento", []),
                "titulos": lp.titulos,
                "objetivo_aprendizagem": lp.objetivo_aprendizagem,
            }
            for lp in _normalize_lines(p)
        ]
    return out

# ---------------------------------------------------------------------
# Renders “placeholder” (substitua pelos reais)
# ---------------------------------------------------------------------
def _render_docx(payload: ExportPayload, linhas: List[LinhaPlano]) -> bytes:
    buf = io.BytesIO()
    buf.write(b"PK\x03\x04")
    return buf.getvalue()

def _render_pdf(payload: ExportPayload, linhas: List[LinhaPlano]) -> bytes:
    return b"%PDF-FAKE\n%..."

def _render_pptx(payload: ExportPayload, linhas: List[LinhaPlano]) -> bytes:
    return b"PK\x03\x04"

def _render_xlsx(payload: ExportPayload, linhas: List[LinhaPlano]) -> bytes:
    return b"PK\x03\x04"

def _split_multiline(val: Optional[str]) -> list[str]:
    """Divide por quebra de linha e ';' (não dividir por '/')."""
    if not val:
        return []
    # normaliza quebras e separa
    raw = str(val).replace("\r\n", "\n")
    parts = re.split(r"[\n;]+", raw)
    # limpa espaços e ignora vazios
    return [re.sub(r"\s+", " ", p).strip() for p in parts if str(p).strip()]
# PATCH-ETAPA: garante que a linha pertence à etapa pedida
def _row_matches_etapa(row: dict, etapa: str) -> bool:
    """
    Retorna True se:
    - não foi pedida etapa (etapa vazia) OU
    - a linha não tem campo de etapa OU
    - a etapa da linha é igual à etapa pedida (case/acento-insensitive)
    """
    e_req = (etapa or "").strip().casefold()
    if not e_req:
        return True  # não filtra por etapa

    e_row = (
        str(row.get("Etapa")
            or row.get("ETAPA")
            or row.get("etapa")
            or "").strip().casefold()
    )

    if not e_row:
        # linha não tem etapa → deixamos passar (seus JSONs antigos)
        return True

    return e_row == e_req


def _normalize_objetivos(val: str | list[str]) -> list[str]:
    """
    Normalização leve: quebra por linha/';', remove bullets/pontuação solta,
    trim, colapsa espaços e dedup preservando ordem — sem reescrever conteúdo.
    """
    txt = "\n".join(val) if isinstance(val, list) else (val or "")
    txt = txt.replace("\r\n", "\n")
    # Quebra por linha e por ';'
    chunks = []
    for line in txt.split("\n"):
        line = re.sub(r'^\s*([*•\-–—]\s*)+', '', line)  # bullets no início
        parts = [p.strip() for p in line.split(";")]
        chunks.extend(parts)
    out, seen = [], set()
    for s in chunks:
        s = re.sub(r'\s+', ' ', s).strip()
        s = re.sub(r'\s*([.,;:])\s*$', '', s)  # pontuação solta no fim
        key = unicodedata.normalize("NFD", s)
        key = "".join(ch for ch in key if unicodedata.category(ch) != "Mn").casefold()
        if s and key not in seen:
            seen.add(key); out.append(s)
    return out


# Chaves por tipo
PL_KEYS = ["Práticas de linguagem","PRÁTICAS DE LINGUAGEM","Práticas de Linguagem","Praticas de linguagem","PRATICAS DE LINGUAGEM"]
UT_KEYS = ["Unidade Temática","Unidade temática","UNIDADE TEMÁTICA","Unidades Temáticas","UNIDADES TEMÁTICAS","UNIDADES TEMATICAS"]
TITLE_KEYS = [
    "Título da aula","TÍTULO DA AULA","Titulo da aula",
    "Título","TÍTULO","Titulo",
    "TÍTULO "  # algumas planilhas vêm com espaço à direita
]

OBJ_KEYS = ["OBJETO DO CONHECIMENTO","OBJETOS DO CONHECIMENTO","Objeto do conhecimento","Objeto de Conhecimento","Objetos de conhecimento"]

# Objetivos – prioriza “OBJETIVOS”
OBJ_APR_PRIMARY  = ["Objetivos de Aprendizagem","OBJETIVOS DE APRENDIZAGEM","Objetivo de Aprendizagem"]
OBJ_APR_FALLBACK = ["OBJETIVOS","Objetivos","Objetivos da Aula","OBJETIVOS DA AULA",
                    "Objetivos de Ensino","OBJETIVOS DE ENSINO"]



def _is_lp(disciplina: str) -> bool:
    d = _norm_key(disciplina)
    # cobre variações com/sem acento/underline
    return "lingua portuguesa" in d or "língua portuguesa" in d

def _tema_source_for(etapa: str, disciplina: str) -> list[str]:
    """
    Define de onde vem o TEMA do contexto:
    - LP em fundamental_I/II => 'Práticas de linguagem'
    - Outros (inclui Médio)  => 'Unidade Temática'
    - Fallback               => Título
    """
    e = (etapa or "").strip()
    if _is_lp(disciplina) and e in ("fundamental_I","fundamental_II"):
        return PL_KEYS + TITLE_KEYS  # tenta PL; se não existir, cai para título
    # demais casos (Médio continua 'Unidade Temática'); fallback por título
    return UT_KEYS + TITLE_KEYS



def _row_objetivos_aprendizagem(row: dict) -> list[str]:
    """Extrai objetivos (sempre priorizando campo 'OBJETIVOS', com fallback)."""
    val = _get_field_ci(row, OBJ_APR_PRIMARY) or _get_field_ci(row, OBJ_APR_FALLBACK)
    return _unique_preserve_order(_split_multiline(val))

OBJ_FALLBACK_KEYS = [
    "Objetos de Conhecimento","Objetos de conhecimento","Objeto de Conhecimento",
    "Objeto do Conhecimento","OBJETO DO CONHECIMENTO","Objetos","Objeto"
]

CONT_FALLBACK_KEYS = ["CONTEÚDO","Conteúdo","Conteudo","CONTEUDO"]

def _int_aula(val) -> Optional[int]:
    """Converte campo 'aula' em inteiro quando possível (tolerante)."""
    if val is None:
        return None
    s = str(val).strip()
    m = re.match(r"^\s*(\d+)", s)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    try:
        return int(float(s.replace(",", ".").split()[0]))
    except Exception:
        return None


def _row_aulas_by_context(row: dict, etapa: str, disciplina: str) -> list[int]:
    """
    Extrai números de aula (ints) usando chaves do field_map e fallback.
    Deduplica e ordena.
    """
    keys = keys_from_field_map(FIELD_MAP, "aula", etapa, disciplina)
    if not keys:
        keys = AULA_KEYS

    # pega o primeiro campo encontrado (ou todos, se preferir)
    val = _get_field_ci(row, keys)
    if val is None:
        return []

    nums: list[int] = []
    # pode vir lista, string única, etc.
    if isinstance(val, list):
        for v in val:
            n = _int_aula(v)
            if isinstance(n, int):
                nums.append(n)
    else:
        # se vier "1; 3; 7"
        for tok in re.split(r"[;\n,|]+", str(val)):
            n = _int_aula(tok)
            if isinstance(n, int):
                nums.append(n)

    # dedupe + ordena
    uniq = sorted({n for n in nums if isinstance(n, int)})
    return uniq


# [PATCH-FIELD-MAP: HABS]  <<< ADICIONAR
HAB_FALLBACK_KEYS = [
    "Habilidade","HABILIDADE","Habilidades","BNCC","BNCC Código","Código","Códigos",
    "RA","RA(s)","Cód. BNCC","Código BNCC"
]

def _row_habilidades_by_context(row: dict, etapa: str, disciplina: str) -> list[str]:
    """
    Extrai habilidades da linha, priorizando o 'field_map' e retornando
    preferencialmente os CÓDIGOS BNCC (EF..., EM..., etc.).
    Se não houver nenhum código, retorna descrições deduplicadas.
    """
    # 1) tenta chaves do field_map
    keys = keys_from_field_map(FIELD_MAP, "habilidade", etapa, disciplina)
    if not keys:
        keys = HAB_KEYS  # fallback curtas
    raw = _get_field_ci(row, keys) or ""

    # quebra multiline: ; , | \n
    parts = _split_multiline(raw)

    codes, descs = [], []
    seen_codes = set()
    seen_descs = set()

    for p in parts:
        s = str(p).strip()
        if not s:
            continue
        m = CODIGO_RE.search(s.upper())
        if m:
            code = m.group(0).upper()
            if code not in seen_codes:
                seen_codes.add(code)
                codes.append(code)
        else:
            k = _norm_key(s)
            if k and k not in seen_descs:
                seen_descs.add(k)
                descs.append(s)

    # Regra: se houver QUALQUER código, retorna só códigos; senão, descrições.
    return codes if codes else descs




def _slug_user(s: Optional[str]) -> str:
    s = str(s or "anon").strip().lower()
    # só letras, números, hífen e underscore
    s = re.sub(r"[^\w\-]+", "-", s, flags=re.UNICODE)
    return s[:64] or "anon"


# PATCH-START helpers

def _parse_list_pipepipe(value: Optional[str]) -> list[str]:
    """Divide 'A||B||C' em lista; limpa vazios/espacos."""
    parts = [p.strip() for p in str(value or "").split("||")]
    return [p for p in parts if p]

def _pick_first_query(params, keys: list[str]) -> Optional[str]:
    """Retorna o primeiro valor não-vazio entre vários aliases de querystring."""
    for k in keys:
        v = params.get(k)
        if v:
            return v
    return None


# no topo do arquivo já existe:
# from utils_norm import norm_text, is_oc_key, OC_CANON_KEY

def _canonize_record_keys(rec: Dict[str, Any]) -> Dict[str, Any]:
    """
    Usa utils_norm para garantir que qualquer forma de 'Objeto(s) do Conhecimento)'
    apareça em rec[OC_CANON_KEY] (tipicamente 'objeto').
    """
    if not isinstance(rec, dict):
        return rec
    if OC_CANON_KEY in rec and str(rec[OC_CANON_KEY]).strip():
        return rec

    out = dict(rec)
    for k in list(rec.keys()):
        if is_oc_key(k):  # ← robusto (singular/plural/typos)
            out[OC_CANON_KEY] = rec[k]
            break
    return out


def _canonize_data_rows(rows: list) -> list:
    return [_canonize_record_keys(r) for r in (rows or [])]

# PATCH-END helpers
DEBUG_ERRORS = True  # deixe True até estabilizar



def _raise_422(route: str, e: Exception):
    if DEBUG_ERRORS:
        raise HTTPException(
            status_code=422,
            detail=f"Falha ao processar {route}: {e.__class__.__name__}: {e}"
        )
    raise HTTPException(status_code=422, detail=f"Falha ao processar {route}")


@app.middleware("http")
async def simple_cors(request: Request, call_next):
    # Pré-flight
    if request.method == "OPTIONS":
        return Response(status_code=204, headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "*",
            "Access-Control-Allow-Headers": "*",
            "Access-Control-Max-Age": "86400",
        })

    # Requisições normais
    try:
        resp = await call_next(request)
    except Exception as e:
        from fastapi.responses import JSONResponse
        resp = JSONResponse(status_code=500, content={"detail": str(e)})

    # Injeta CORS em toda resposta
    resp.headers["Access-Control-Allow-Origin"] = "*"
    resp.headers["Access-Control-Allow-Methods"] = "*"
    resp.headers["Access-Control-Allow-Headers"] = "*"
    return resp



async def add_cors_headers(request: Request, call_next):
    resp = await call_next(request)
    resp.headers.setdefault("Access-Control-Allow-Origin", "*")
    resp.headers.setdefault("Access-Control-Allow-Methods", "*")
    resp.headers.setdefault("Access-Control-Allow-Headers", "*")
    return resp

@app.options("/{path:path}")
def options_cors(path: str):
    return Response(status_code=204, headers={
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "*",
        "Access-Control-Allow-Headers": "*",
    })


@app.get("/api/ping")
def ping():
    return {"ok": True} 
# ---------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------
def _safe_json(p: Path):
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        return None
        

def _uniq(seq: List[Any]) -> List[Any]:
    seen, out = set(), []
    for x in seq:
        if x not in seen:
            out.append(x)
            seen.add(x)
    return out

def _slug(s: Optional[str]) -> str:
    if not s: return ""
    t = re.sub(r"[^\w\s-]", "", str(s), flags=re.UNICODE)
    t = t.strip().lower()
    t = re.sub(r"[\s_-]+", "-", t)
    return t[:80]


def _normalize(s: str) -> str:
    if not s:
        return ""
    # remove acentos e converte para lowercase
    s = unicodedata.normalize("NFD", s)
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    return s.casefold().strip()


def _flatten_lines(val) -> List[str]:
    if val is None: return []
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    s = str(val).strip()
    return [s] if s else []

def _titles_from_payload(payload: Dict[str, Any]) -> List[str]:
    """
    Extrai títulos da aula do payload do front, aceitando várias convenções:
    - Strings, listas de strings, listas de objetos {label|titulo|title|nome|name: ...}
    - Vários nomes de campo possíveis vindos do front.
    """
    if not isinstance(payload, dict):
        return []

    # possíveis chaves vindas do front
    candidates = [
        "titulos_da_aula", "titulos", "titulo",
        "titulosAula", "titulos_aula",
        "titulosSelecionados", "titulos_selecionados",
        "partes", "partes_tema", "tema_partes",
        "selecionados_titulo", "selecionadosTitulos",
        "selTitulo"  # em alguns fronts salvam direto com o id do select
    ]

    raw = None
    for k in candidates:
        if k in payload and payload[k] not in (None, "", []):
            raw = payload[k]
            break

    out: List[str] = []
    def push(v):
        s = str(v or "").strip()
        if s and s not in out:
            out.append(s)

    if raw is None:
        return out

    # string única (pode vir separada por \n ou "||")
    if isinstance(raw, str):
        parts = [p.strip() for p in re.split(r"\n+|\|\|", raw) if str(p).strip()]
        for p in parts: push(p)
        return out

    # lista de strings
    if isinstance(raw, list) and all(isinstance(x, str) or x is None for x in raw):
        for x in raw:
            push(x)
        return out

    # lista de objetos (label/title/titulo/nome/name)
    if isinstance(raw, list) and any(isinstance(x, dict) for x in raw):
        for obj in raw:
            if not isinstance(obj, dict):
                continue
            val = (
                obj.get("label") or obj.get("titulo") or obj.get("title") or
                obj.get("nome")  or obj.get("name")   or obj.get("text")
            )
            if val:
                push(val)
        return out

    # objeto único (ex.: {label: "..."} )
    if isinstance(raw, dict):
        val = (
            raw.get("label") or raw.get("titulo") or raw.get("title") or
            raw.get("nome")  or raw.get("name")   or raw.get("text")
        )
        if val:
            push(val)
    return out


    def _norm_txt(s: str) -> str:
        s = unicodedata.normalize("NFD", str(s or ""))
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        return re.sub(r"\s+", " ", s).strip().casefold()

    def _clean_hab_line(s: str) -> str:
        # remove bullet/numeração + código BNCC no início; depois normaliza
        s = re.sub(r'^\s*[•\-\*\u2022]\s*', '', str(s or ""))
        s = CODIGO_RE.sub("", s)
        return _norm_txt(s)

    def _split_habilidades(texto: str) -> list[str]:
        """
        Reconstrói uma lista de habilidades a partir de um bloco de texto possivelmente quebrado.
        - Inicia nova habilidade se a linha começa com código BNCC/RA, bullet (•, -, *) ou numeração.
        - Caso contrário, considera continuação da habilidade anterior.
        """
        if not texto:
            return []

        lines = [ln.strip() for ln in str(texto).replace('\r\n', '\n').split('\n')]

        itens: list[str] = []
        buf: list[str] = []

    def flush():
        if buf:
            s = ' '.join(buf).strip()
            s = re.sub(r'\s{2,}', ' ', s)  # comprime espaços múltiplos
            if s:
                itens.append(s)
            buf.clear()

    for ln in lines:
        if not ln:
            continue

        starts_with_code   = bool(CODIGO_RE.match(ln))
        starts_with_bullet = bool(re.match(r'^\s*[•\-\*\u2022]\s*', ln))
        starts_with_number = bool(re.match(r'^\s*\(?\d+[\).\-\s]\s*', ln))

        if starts_with_code or starts_with_bullet or starts_with_number:
            flush()
            ln = re.sub(r'^\s*[•\-\*\u2022]\s*', '', ln)
            ln = re.sub(r'^\s*\(?\d+[\).\-\s]\s*', '', ln)
            buf.append(ln)
        else:
            if not buf:
                buf.append(ln)
            else:
                buf.append(ln)

    flush()

    seen = set(); out = []
    for s in itens:
        if s not in seen:
            seen.add(s)
            out.append(s)
    return out

def _find_arquivo(etapa: str, disciplina: str) -> str:
    etapa_norm = _normalize(etapa)
    disc_norm  = _normalize(disciplina)
    man = _load_manifest()
    for it in man:
        if _normalize(it.get("etapa")) == etapa_norm and _normalize(it.get("disciplina")) == disc_norm:
            return it.get("arquivo") or ""
    raise HTTPException(status_code=404, detail=f"Disciplina não encontrada no manifest: {disciplina}")
TEMA_KEYS = [
    # básicas + fallback por título da aula
    "Unidade Temática","Unidade temática","UNIDADE TEMÁTICA","Unidades Temáticas",
    "TEMA","Tema","tema","Práticas de linguagem", "PRÁTICAS DE LINGUAGEM", 
    "Práticas de Linguagem","Linguagem", "linguagem", "LINGUAGEM", "linguagens","Linguagens","LINGUAGENS"
    
]
HAB_KEYS  = [
    # básicas
    "HABILIDADE","Habilidades","Habilidade","habilidade",
    "CÓDIGO HABILIDADE","CÓDIGO DA HABILIDADE","Código","codigo","CÓDIGO",
    "Descrição da Habilidade","Descricao da Habilidade","descrição","Descricao","Código BNCC","Código(s) BNCC","BNCC","RA","RA(s)","Código","Códigos",
    # técnico (competência/RA)
    "Competência","COMPETÊNCIA","Competências","COMPETÊNCIAS","Unidade de Competência",
    "Resultados de aprendizagem","RA"
]
# Colunas candidatas para "Objeto do Conhecimento" (conteúdo)
# Colunas candidatas para CONTEÚDO (somente conteúdo/tópico/assunto)
CONTEUDO_KEYS = [
    "Conteúdo","Conteúdos","CONTEÚDO","CONTEÚDOS",
    "Tópico","Tópicos","Assunto"
]


# Colunas candidatas para objetivos
OBJ_APR_KEYS = [
    "Objetivos de Aprendizagem","OBJETIVOS DE APRENDIZAGEM","Objetivo de Aprendizagem",
    "Objetivos","OBJETIVOS","Objetivos da Aula","OBJETIVOS DA AULA",
    "Objetivos de Ensino","OBJETIVOS DE ENSINO"
]
OBJ_ESP_KEYS = [
    "Objetivos Específicos","OBJETIVOS ESPECÍFICOS","Objetivos específicos",
    "Resultados Esperados","RESULTADOS ESPERADOS","Resultados",
    "Objetivos Detalhados","OBJETIVOS DETALHADOS"
]

# --- helper: normaliza leitura de um manifest (lista de itens) ---
def _collect_manifest_items_from_file(p: Path) -> List[Dict[str, Any]]:
    data = _safe_json(p)
    # alguns manifests podem vir como {"itens": [...]}
    if isinstance(data, dict) and "itens" in data:
        data = data.get("itens")
    if not isinstance(data, list):
        return []
    items = []
    for it in data:
        etapa = str(it.get("etapa", "")).strip()
        disciplina = str(it.get("disciplina", "")).strip()
        arquivo = str(it.get("arquivo", "")).strip()
        # se o arquivo vier relativo, garante posix ("/") e normaliza
        if arquivo:
            arquivo_rel = Path(arquivo)
            if arquivo_rel.is_absolute():
                try:
                    arquivo_rel = arquivo_rel.relative_to(DATA_DIR)
                except Exception:
                    # mantém como veio, mas forçando barra '/'
                    arquivo_rel = Path(arquivo_rel.name)
            arquivo = arquivo_rel.as_posix()
        elif etapa and disciplina:
            # fallback se não veio "arquivo" no manifest
            arquivo = f"{etapa}/{disciplina}.json"
        items.append({"etapa": etapa, "disciplina": disciplina, "arquivo": arquivo})
    return items

# --- novo loader tolerante ---
def _load_manifest() -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []

    # 1) manifest.json combinado (se existir)
    comb = DATA_DIR / "manifest.json"
    if comb.exists():
        items.extend(_collect_manifest_items_from_file(comb))

    # 2) qualquer manifest_*.json (exceto o combinado acima)
    for p in DATA_DIR.glob("manifest*.json"):
        if p.name == "manifest.json":
            continue
        items.extend(_collect_manifest_items_from_file(p))

    # 3) fallback: fazer SCAN das subpastas (fundamental_I, medio, etc.)
    #    e montar linhas a partir dos arquivos .json encontrados
    for etapa_dir in DATA_DIR.iterdir():
        if etapa_dir.is_dir():
            etapa = etapa_dir.name
            for f in etapa_dir.glob("*.json"):
                items.append({
                    "etapa": etapa,
                    "disciplina": f.stem,
                    "arquivo": f.relative_to(DATA_DIR).as_posix()
                })

    # deduplica por (etapa, disciplina, arquivo) e garante chaves
    seen = set()
    final: List[Dict[str, Any]] = []
    for it in items:
        etapa = it.get("etapa", "")
        disc = it.get("disciplina", "")
        arq = it.get("arquivo", "")
        if not arq and etapa and disc:
            arq = f"{etapa}/{disc}.json"
        key = (etapa, disc, arq)
        if key in seen:
            continue
        seen.add(key)
        final.append({"etapa": etapa, "disciplina": disc, "arquivo": arq})
    return final

def _flex_bool(v) -> bool:
    """
    Converte strings/valores comuns de query em boolean:
    aceita 1/true/t/yes/y/on/sim (case-insensitive).
    """
    return str(v).strip().lower() in {"1", "true", "t", "yes", "y", "on", "sim"}

def _collect_unique_by_keys(rows: list[dict], keys: list[str]) -> list[str]:
    """Coleta valores únicos (lista ou string) olhando várias chaves possíveis."""
    seen = set()
    out  = []
    for r in rows or []:
        for k in keys:
            if k in r and str(r[k]).strip():
                v = r[k]
                if isinstance(v, (list, tuple, set)):
                    for x in v:
                        s = re.sub(r"\s+", " ", str(x)).strip()
                        kx = norm_key_ci(s)
                        if s and kx not in seen:
                            seen.add(kx); out.append(s)
                else:
                    s = re.sub(r"\s+", " ", str(v)).strip()
                    kx = norm_key_ci(s)
                    if s and kx not in seen:
                        seen.add(kx); out.append(s)
                break
    return {"temas": _finalize_list(out)}



# ---------------------------------------------------------------------
# Rotas – Básicas/Debug
# ---------------------------------------------------------------------
@app.get("/api/health")
def health():  # simples texto para acesso direto
    return "ok"

@app.get("/api/version")
def version():
    return {"name": APP_NAME, "version": APP_VERSION}

@app.get("/api/debug/manifest-info")
def manifest_info():
    man = _load_manifest()
    return {
        "qtd_itens": len(man),
        "etapas": sorted({m.get("etapa", "") for m in man if m.get("etapa")})
    }

@app.get("/api/manifest")
def get_manifest(etapa: Optional[str] = Query(None, description="Filtra por etapa")):
    man = _load_manifest()
    if etapa:
        man = [m for m in man if m.get("etapa") == etapa]
    return man

@app.get("/api/debug/disciplinas")
def debug_disciplinas(etapa: str = Query(..., description="ex.: medio, fundamental_I, fundamental_II")):
    etapa_norm = normalize_etapa(etapa)
    rows = _load_rows_canon(etapa_norm, None, None)

    disc_vals = []
    for r in rows:
        d = _get_field_ci(r, ALIASES.get("disciplina", ["disciplina","Disciplina"]))
        if d and str(d).strip():
            disc_vals.append(str(d).strip())
    cnt = Counter(disc_vals)
    # retorna ordenado por frequência
    out = [{"disciplina": k, "linhas": v} for k, v in cnt.most_common()]
    return {"etapa": etapa_norm, "disciplinas": out}
# ---------------------------------------------------------------------
# Rotas – Navegação por etapa/arquivo
# ---------------------------------------------------------------------
@app.get("/api/disciplinas")
def get_disciplinas(etapa: str = Query(..., description="ex.: medio")):
    man = _load_manifest()
    # lista bruta da sua fonte
    raw = [m.get("disciplina", "") for m in man if m.get("etapa") == etapa and m.get("disciplina")]

    # dedupe por chave normalizada e preferir sem underline
    seen = {}  # key normalizada -> rótulo “bonito” (com espaços)
    for nome in raw:
        key = _norm_key(nome)
        if IGNORE_PAT.match(key):
            continue
        label = nome.replace("_", " ")
        cur = seen.get(key)
        if cur is None or ("_" in cur and "_" not in nome):
            seen[key] = label

    disciplinas = sorted(seen.values(), key=lambda x: unicodedata.normalize("NFKD", x).casefold())
    return {"disciplinas": disciplinas}

@app.get("/api/arquivo-da-disciplina")
def get_arquivo_da_disciplina(
    etapa: str = Query(...),
    disciplina: str = Query(...)
):
    key = _norm_key(disciplina)

    # 1) tenta casar pelo manifest usando nome normalizado
    try:
        man = _load_manifest()
    except Exception:
        man = []

    for item in man:
        if item.get("etapa") == etapa and _norm_key(item.get("disciplina", "")) == key:
            arq = item.get("arquivo")
            if arq:
                return {"arquivo": arq}

    # 2) fallback: tenta achar arquivo com/sem underline e sem acentos
    base = (
        DATA_DIR / etapa
        if 'DATA_DIR' in globals()
        else Path(__file__).resolve().parent / "app" / "data" / etapa
    )

    def _no_acc(s: str) -> str:
        s = unicodedata.normalize("NFKD", s or "")
        return "".join(ch for ch in s if not unicodedata.combining(ch))

    candidatos = []
    for nome in [disciplina]:
        if not nome:
            continue
        formas = {
            nome,
            nome.replace(" ", "_"),
            nome.replace("_", " "),
            _no_acc(nome),
            _no_acc(nome).replace(" ", "_"),
            _no_acc(nome).replace("_", " "),
        }
        candidatos.extend(sorted(formas, key=len, reverse=True))  # tenta as mais longas primeiro

    # este loop estava com identação errada no seu código
    for cand in candidatos:
        p = base / f"{cand}.json"
        if p.exists():
            return {"arquivo": f"{etapa}/{cand}.json"}

    raise HTTPException(
        status_code=404,
        detail=f"Arquivo não encontrado para {disciplina} ({etapa})."
    )




# ---------------------------------------------------------------------
# Rotas – Dados (temas / habilidades) – via arquivo
# ---------------------------------------------------------------------
@app.get("/api/dados/temas")
def dados_temas(
    arquivo: Optional[str] = Query(None),
    etapa:   Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    conteudo: Optional[str] = Query(None),
    contains: Optional[str] = Query(None),
):
    try:
        e_norm = normalize_etapa(etapa) if etapa else None
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        tema_list     = split_multi(tema)
        obj_list      = split_multi(objeto)
        titulo_list   = split_multi(titulo)
        conteudo_list = split_multi(conteudo)

        rows = _load_rows_canon(e_norm, d_norm, arquivo) or []

        # Sem filtros finos → devolve todos os temas
        if not any([tema_list, obj_list, titulo_list, conteudo_list]):
            todos = []
            for r in rows:
                try:
                    todos.extend(row_temas(r, FIELD_MAP, e_norm, d_norm) or [])
                except Exception as _e:
                    # log defensivo por linha, não falha o endpoint
                    print("[/api/dados/temas] WARN row_temas:", repr(_e))
                    continue
            return {"temas": finalize_list(todos)}

        # Com filtros: use seu coletor unificado (se existir), senão filtre localmente
        try:
            result = filter_and_collect(
                data=rows,
                field_map=FIELD_MAP,
                target_alias_key="tema",
                tema=tema, objeto=objeto, titulo=titulo, conteudo=conteudo,
                contains=contains_b, etapa=e_norm, disciplina=d_norm,
            )
            return {"temas": finalize_list(result)}
        except NameError:
            # Fallback: filtro local leve se filter_and_collect não existir
            wanted_tema   = {s.casefold() for s in tema_list}
            wanted_obj    = {s.casefold() for s in obj_list}
            wanted_titulo = {s.casefold() for s in titulo_list}
            wanted_cont   = {s.casefold() for s in conteudo_list}

            out = []
            for r in rows:
                temas   = [t.casefold() for t in (row_temas(r, FIELD_MAP, e_norm, d_norm) or [])]
                objs    = [o.casefold() for o in (row_objetos(r, FIELD_MAP, e_norm, d_norm) or [])]
                titulos = [t.casefold() for t in (row_titles(r, FIELD_MAP, e_norm, d_norm) or [])]
                # match simples (contains_b ignora aqui)
                if wanted_tema   and not (set(temas)   & wanted_tema):   continue
                if wanted_obj    and not (set(objs)    & wanted_obj):    continue
                if wanted_titulo and not (set(titulos) & wanted_titulo): continue
                out.extend(temas)
            return {"temas": finalize_list(out)}
    except Exception as e:
        # LOGA parâmetros para depuração
        print("[/api/dados/temas] ERRO:", repr(e))
        print("   params:", dict(
            arquivo=arquivo, etapa=etapa, disciplina=disciplina,
            tema=tema, objeto=objeto, titulo=titulo, conteudo=conteudo, contains=contains
        ))
        raise HTTPException(status_code=422, detail=f"Falha ao processar /api/dados/temas: {e.__class__.__name__}: {e}")





@app.get("/api/dados/objetos")
def dados_objetos(
    arquivo: Optional[str] = Query(None),
    etapa:   Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),    # filtro opcional
    titulo: Optional[str] = Query(None),
    conteudo: Optional[str] = Query(None),
    contains: Optional[str] = Query(None),
):
    try:
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        if not d_norm and not arquivo:
            return {"objetos": []}

        tema_list     = split_multi(tema)
        obj_list      = split_multi(objeto)
        titulo_list   = split_multi(titulo)
        conteudo_list = split_multi(conteudo)

        rows = _load_rows_canon(e_norm, d_norm, arquivo)

        if not any([tema_list, obj_list, titulo_list, conteudo_list]):
            todos = []
            for r in rows:
                todos.extend(row_objetos(r, FIELD_MAP, e_norm, d_norm) or [])
            return {"objetos": finalize_list(todos)}

        result = filter_and_collect(
            data=rows,
            field_map=FIELD_MAP,
            target_alias_key="objeto",
            tema=tema, objeto=objeto, titulo=titulo, conteudo=conteudo,
            contains=contains_b, etapa=e_norm, disciplina=d_norm,
        )
        return {"objetos": finalize_list(result)}
    except Exception as e:
        print("[/api/dados/objetos] ERRO:", repr(e))
        raise HTTPException(status_code=422, detail=f"Falha ao processar /api/dados/objetos: {e.__class__.__name__}: {e}")

@app.get("/api/dados/titulos")
def dados_titulos(
    arquivo: Optional[str] = Query(None),
    etapa:   Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),  # filtro opcional de título
    conteudo: Optional[str] = Query(None),
    aula: Optional[str]     = Query(None),
    contains: Optional[str] = Query(None),
):
    try:
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        if not d_norm and not arquivo:
            return {"titulos": []}

        tema_list     = split_multi(tema)
        obj_list      = split_multi(objeto)
        titulo_list   = split_multi(titulo)
        conteudo_list = split_multi(conteudo)
        aula_list     = split_multi(aula)

        rows = _load_rows_canon(e_norm, d_norm, arquivo)

        if not any([tema_list, obj_list, titulo_list, conteudo_list, aula_list]):
            todos = []
            for r in rows:
                todos.extend(row_titles(r, FIELD_MAP, e_norm, d_norm) or [])
            return {"titulos": finalize_list(todos)}

        result = filter_and_collect(
            data=rows,
            field_map=FIELD_MAP,
            target_alias_key="titulo",
            tema=tema, objeto=objeto, titulo=titulo, conteudo=conteudo, aula=aula,
            contains=contains_b, etapa=e_norm, disciplina=d_norm,
        )
        return {"titulos": finalize_list(result)}
    except Exception as e:
        print("[/api/dados/titulos] ERRO:", repr(e))
        raise HTTPException(status_code=422, detail=f"Falha ao processar /api/dados/titulos: {e.__class__.__name__}: {e}")



@app.get("/api/dados/conteudos")
def dados_conteudos(
    arquivo: Optional[str] = Query(None),
    etapa:   Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    aula: Optional[str]   = Query(None),
    contains: Optional[str] = Query(None),
):
    """
    Retorna CONTEÚDOS filtrados.
    - Aceita chamadas só com etapa+disciplina (sem filtros de contexto).
    - Quando nenhum filtro de contexto vem, devolve TODOS os conteúdos da disciplina.
    """
    try:
        # 🔹 Normalização básica de etapa/disciplina
        e_norm = normalize_etapa(etapa) if etapa else None
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        # Se não há disciplina normalizada nem arquivo explícito, não temos de onde ler
        if not d_norm and not arquivo:
            return {"conteudos": []}

        # 🔹 Blindagem de filtros: sempre listas
        tema_list   = split_multi(tema)   or []
        obj_list    = split_multi(objeto) or []
        titulo_list = split_multi(titulo) or []
        aula_list   = split_multi(aula)   or []

        # 🔹 Carrega as linhas (wrapper compatível)
        rows = _load_rows_canon(e_norm, d_norm, arquivo)

        # 🔹 Sem nenhum filtro de contexto -> devolve TODOS os conteúdos
        if not any([tema_list, obj_list, titulo_list, aula_list]):
            todos: list[str] = []
            for r in rows:
                # Usa a função canônica de extração de conteúdos
                todos.extend(_row_conteudos(r, FIELD_MAP, e_norm, d_norm) or [])

            return {"conteudos": finalize_list(todos)}

        # 🔹 Com filtros -> usa filter_and_collect
        result = filter_and_collect(
            data=rows,
            field_map=FIELD_MAP,
            target_alias_key="conteudo",
            tema=tema_list,
            objeto=obj_list,
            titulo=titulo_list,
            aula=aula_list,
            contains=contains_b,
            etapa=e_norm,
            disciplina=d_norm,
        )
        return {"conteudos": finalize_list(result)}

    except HTTPException:
        # Se alguma parte do código já levantou HTTPException explícita, repassa
        raise
    except Exception as e:
        # Aqui, qualquer erro interno vira 500 (mais adequado que 422)
        print("[/api/dados/conteudos] ERRO:", repr(e))
        raise HTTPException(
            status_code=500,
            detail=f"Falha ao processar /api/dados/conteudos: {e.__class__.__name__}: {e}",
        )



# ---------------------------------------------------------------------
# /api/dados/habilidades  (VERSÃO NOVA, PADRONIZADA)
# ---------------------------------------------------------------------
@app.get("/api/dados/habilidades")
def dados_habilidades(
    etapa: str = Query(...),
    disciplina: str = Query(...),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    conteudo: Optional[str] = Query(None),
    aula: Optional[str] = Query(None),
    contains: bool = Query(False),
    only_codes: bool = Query(False),
):
    """
    Retorna habilidades filtradas por contexto (tema/objeto/título/conteúdo/aula).

    Nesta versão:
    - Não usa filter_and_collect_habilidades (que parece estar quebrando).
    - Usa match_row + _row_habilidades diretamente, com try/except por linha.
    - Nunca estoura 500 por causa de uma linha problemática.
    """

    # ===== DEBUG: entrada bruta =====
    print("[dados_habilidades] IN (parâmetros brutos):", {
        "etapa": etapa,
        "disciplina": disciplina,
        "tema": tema,
        "objeto": objeto,
        "titulo": titulo,
        "conteudo": conteudo,
        "aula": aula,
        "contains": contains,
        "only_codes": only_codes,
    })

    try:
        # 1) Normaliza etapa/disciplina
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, disciplina)

        print("[dados_habilidades] normalizado:", {
            "etapa_norm": e_norm,
            "disciplina_norm": d_norm,
        })

        # 2) Carrega linhas do dataset da disciplina
        try:
            rows = load_rows_for(e_norm, d_norm)
        except HTTPException as http_exc:
            if http_exc.status_code == 404:
                print(
                    f"[dados_habilidades] dataset 404 para {e_norm}/{d_norm} "
                    "→ retornando lista vazia de habilidades"
                )
                return []
            print(
                "[dados_habilidades] HTTPException ao carregar rows:",
                http_exc.status_code, http_exc.detail
            )
            raise

        total_rows = len(rows)
        print(f"[dados_habilidades] rows carregadas = {total_rows}")
        if total_rows > 0:
            print("[dados_habilidades] amostra linha[0] chaves =",
                  list(rows[0].keys()))
        if total_rows > 1:
            print("[dados_habilidades] amostra linha[1] chaves =",
                  list(rows[1].keys()))

        # 3) Monta contexto de filtro (tema/objeto/título/conteúdo/aula)
        ctx = {
            "tema": tema,
            "objeto": objeto,
            "titulo": titulo,
            "conteudo": conteudo,
            "aula": aula,
        }
        use_ctx = any(
            (v is not None) and (str(v).strip() != "")
            for v in ctx.values()
        )

        print("[dados_habilidades] use_ctx =", use_ctx, "ctx =", ctx)

        todas_habs: List[str] = []

        # IMPORTANTE:
        # - assumindo que você já tem FIELD_MAP carregado em escopo global
        #   a partir do field_map.json (se o nome for outro, ajuste aqui).
        global FIELD_MAP

        for idx, row in enumerate(rows):
            try:
                # 3a) Se tiver contexto, filtra a linha antes
                if use_ctx:
                    if not match_row(
                        row=row,
                        etapa=e_norm,
                        disciplina=d_norm,
                        tema=tema,
                        objeto=objeto,
                        titulo=titulo,
                        conteudo=conteudo,
                        aula=aula,
                        field_map=FIELD_MAP,
                        contains=contains,
                    ):
                        continue

                # 3b) Extrai habilidades da linha
                habs = _row_habilidades(row, FIELD_MAP, e_norm, d_norm)

            except Exception as row_exc:
                # DEBUG por linha problemática (mas sem matar o endpoint)
                print(f"[dados_habilidades] ERRO ao processar linha {idx}: {row_exc}")
                print("[dados_habilidades]   chaves linha:", list(row.keys()))
                continue

            if not habs:
                continue

            for h in habs:
                if not isinstance(h, str):
                    continue
                h_clean = h.strip()
                if not h_clean:
                    continue
                todas_habs.append(h_clean)

        # 4) Dedup preservando ordem
        habilidades = unique_preserve_order(todas_habs)

        # Se o front pedir apenas códigos (only_codes), você pode
        # futuramente aplicar um regex aqui. Por enquanto, devolve tudo.
        if only_codes:
            print("[dados_habilidades] only_codes=True, mas ainda não filtrando só códigos (devolvendo tudo).")

        print(f"[dados_habilidades] total habilidades retornadas = {len(habilidades)}")
        if habilidades:
            print("[dados_habilidades] amostra habilidades[0:3] =",
                  habilidades[:3])

        return habilidades

    except HTTPException:
        # deixa FastAPI tratar HTTPException normalmente
        raise
    except Exception:
        import traceback
        tb = traceback.format_exc()
        print("[ERRO dados_habilidades] Erro inesperado (nível topo):")
        print(tb)
        raise HTTPException(
            status_code=500,
            detail="Erro interno ao processar habilidades."
        )



def _match_list(row_vals: list[str], wanted: set[str], partial: bool) -> bool:
    if not wanted:
        return True
    row_norm = {norm_key_ci(v) for v in (row_vals or []) if str(v).strip()}
    if not row_norm:
        return False
    if row_norm & wanted:
        return True
    if partial:
        for w in wanted:
            for r in row_norm:
                if w in r or r in w:
                    return True
    return False

def _same_ci(a: str, b: str) -> bool:
    try:
        return _norm_ci(a) == _norm_ci(b)
    except Exception:
        return (a or "").strip().casefold() == (b or "").strip().casefold()



def _load_rows_canon(etapa_norm: Optional[str], disc_norm: Optional[str], arquivo: Optional[str]) -> List[Dict[str, Any]]:
    """
    Loader canônico **simplificado**:
    - Hoje delega para load_rows_for(etapa, disciplina), que já está testado e funcionando
      para temas/objetos/títulos/aulas.
    - Ignora 'arquivo' por enquanto (você pode reintroduzir esse comportamento depois,
      de forma controlada).
    """
    if not etapa_norm or not disc_norm:
        return []

    # Se load_rows_for já cuida de normalizar internamente, você pode passar etapa/disciplinas
    # "originais". Se ele espera normalizadas, mantenha e_norm/d_norm como está.
    try:
        return load_rows_for(etapa_norm, disc_norm) or []
    except HTTPException as e:
        # Se houver algum erro específico de dados, você decide se quer propagar ou
        # transformar em lista vazia.
        if e.status_code == 404:
            # Sem dados para essa combinação → devolve vazio
            return []
        raise
    except Exception as e:
        # Em vez de deixar erros estranhos vazarem, loga e devolve vazio (ou re-raise).
        print("[_load_rows_canon] erro inesperado:", repr(e))
        return []



@app.get("/api/dados/aulas")
def dados_aulas(
    etapa: Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None, description="Tema(s) — 'A||B'"),
    contains: bool = Query(False)
) -> List[str]:
    """
    Retorna aulas para o contexto. Fallback: deriva quantidade a partir dos títulos.
    """
    # TODO: se tiver storage próprio, tente carregar aqui:
    # aulas = try_get_aulas_from_store(etapa, disciplina, tema, contains)
    # if aulas: return [str(x).strip() for x in aulas if str(x).strip()]

    # Fallback: derive dos títulos
        # Fallback: derive dos títulos
        # Fallback: derivar quantidade a partir dos TÍTULOS já que não há storage nativo de “aulas”
    try:
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = bool(contains)

        rows = _load_rows_canon(e_norm, d_norm, None)

        # Se não veio tema, usa todos os títulos; se veio, filtra por tema
        tema_list = split_multi(tema)
        if not tema_list:
            todos_titulos = []
            for r in rows:
                todos_titulos.extend(row_titles(r, FIELD_MAP, e_norm, d_norm) or [])
            titulos = finalize_list(todos_titulos)
        else:
            titulos = filter_and_collect(
                data=rows,
                field_map=FIELD_MAP,
                target_alias_key="titulo",
                tema=tema, objeto=None, titulo=None, conteudo=None,
                contains=contains_b, etapa=e_norm, disciplina=d_norm,
            )

        n = len(titulos or [])
        return [f"Aula {i}" for i in range(1, n + 1)] if n > 0 else []
    except Exception as e:
        print("[/api/dados/aulas] ERRO:", repr(e))
        raise HTTPException(status_code=422, detail=f"Falha ao processar /api/dados/aulas: {e.__class__.__name__}: {e}")




# ---------------------------------------------------------------------
# /api/dados/objetivos
# ---------------------------------------------------------------------

@app.get("/api/dados/objetivos")
def dados_objetivos(
    arquivo:    Optional[str] = Query(None),
    etapa:      Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema:       Optional[str] = Query(None),
    objeto:     Optional[str] = Query(None),
    titulo:     Optional[str] = Query(None),
    conteudo:   Optional[str] = Query(None),
    contains:   Optional[str] = Query(None),
):
    """
    Retorna 'Objetivos' (linhas) extraídos das planilhas/JSONs, respeitando:
      - field_map por etapa/disciplina (aliases de 'objetivos')
      - filtros: tema, objeto, título, conteúdo
      - contains: correspondência parcial ('contém') quando =1/true
    """
    try:
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        # sem disciplina e sem arquivo → nada a listar
        if not d_norm and not arquivo:
            return {"objetivos": []}

        # carregar linhas normalizadas do dataset
        rows = _load_rows_canon(e_norm, d_norm, arquivo)

        # se não veio nenhum filtro fino, devolve todos os objetivos
        if not any([tema, objeto, titulo, conteudo]):
            todos = filter_and_collect(
                data=rows,
                field_map=FIELD_MAP,
                target_alias_key="objetivos",
                tema=None, objeto=None, titulo=None, conteudo=None,
                contains=False, etapa=e_norm, disciplina=d_norm,
            )
            return {"objetivos": finalize_list(todos)}

        # com filtros → usa coletor unificado
        result = filter_and_collect(
            data=rows,
            field_map=FIELD_MAP,
            target_alias_key="objetivos",
            tema=tema, objeto=objeto, titulo=titulo, conteudo=conteudo,
            contains=contains_b, etapa=e_norm, disciplina=d_norm,
        )
        return {"objetivos": finalize_list(result)}

    except Exception as e:
        print("[/api/dados/objetivos] ERRO:", repr(e))
        raise HTTPException(
            status_code=422,
            detail=f"Falha ao processar /api/dados/objetivos: {e.__class__.__name__}: {e}",
        )


# ---------------------------------------------------------------------
# Núcleo compartilhado para /api/dados/linhas e /api/dados/conteudos
# ---------------------------------------------------------------------

def _filtrar_linhas_por_contexto(
    etapa: str,
    disciplina: str,
    tema: Optional[str],
    objeto: Optional[str],
    titulo: Optional[str],
    conteudo: Optional[str],
    aula: Optional[str],
    habilidade: Optional[str],
    contains: bool,
):
    """
    Núcleo de filtragem usado por /api/dados/linhas e /api/dados/conteudos.

    Retorna **linhas cruas** do JSON, já filtradas por:
    tema / objeto / título / conteúdo / aula (e futuramente habilidade).
    """
    try:
        # 1) Normalização de etapa/disciplina
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, disciplina)

        # 2) Carrega todas as linhas dessa etapa/disciplina
        rows = load_rows_for(e_norm, d_norm)
        print(f"[dados_core] {e_norm}/{d_norm} → {len(rows)} linhas carregadas")

        # 3) Quebra 'A||B' em listas (usa a mesma split_multi dos outros endpoints)
        tema_list       = split_multi(tema)       or []
        objeto_list     = split_multi(objeto)     or []
        titulo_list     = split_multi(titulo)     or []
        conteudo_list   = split_multi(conteudo)   or []
        aula_list       = split_multi(aula)       or []
        habilidade_list = split_multi(habilidade) or []

        print(
            "[dados_core] filtros: "
            f"tema={tema_list}, objeto={objeto_list}, titulo={titulo_list}, "
            f"conteudo={conteudo_list}, aula={aula_list}, habilidade={habilidade_list}, "
            f"contains={contains}"
        )

        # Se pelo menos um filtro de contexto foi definido
        tem_contexto = any([
            tema_list,
            objeto_list,
            titulo_list,
            conteudo_list,
            aula_list,
            habilidade_list,  # por enquanto só indica que há filtro por habilidade
        ])

        # 4) Sem filtros + contains=False → devolve tudo
        if not tem_contexto and not contains:
            print(f"[dados_core] sem filtros → retornando todas as {len(rows)} linhas")
            return rows  # lista crua de linhas do JSON

        # 5) Com filtros → prepara seleções e usa match_row
        temas_sel     = set(tema_list)     if tema_list     else set()
        objetos_sel   = set(objeto_list)   if objeto_list   else set()
        titulos_sel   = set(titulo_list)   if titulo_list   else set()
        conteudos_sel = set(conteudo_list) if conteudo_list else set()
        aulas_sel     = set(aula_list)     if aula_list     else set()
        # OBS: habilidade_list ainda não é passado pra match_row; podemos
        # incluir depois, se desejado, dentro da própria match_row.

        filtradas: list[dict] = []
        for idx, r in enumerate(rows):
            try:
                ok = match_row(
                    r,
                    FIELD_MAP,          # field_map
                    temas_sel,          # temas_sel
                    objetos_sel,        # objs_sel
                    titulos_sel,        # tits_sel
                    conteudos_sel,      # cont_sel
                    aulas_sel,          # aulas_sel
                    contains,           # contains (bool)
                    etapa=e_norm,
                    disciplina=d_norm,
                )
            except TypeError as te:
                # TypeError em match_row NÃO deve derrubar o endpoint inteiro
                print(
                    f"[dados_core] TypeError em match_row na linha #{idx}: "
                    f"{type(te).__name__}: {te}"
                )
                # pula linha problemática, segue com as demais
                continue
            except Exception as e:
                print(
                    f"[dados_core] ERRO em match_row na linha #{idx}: "
                    f"{type(e).__name__}: {e}"
                )
                # também pula linha em caso de erro genérico
                continue

            if ok:
                filtradas.append(r)

        print(f"[dados_core] com filtros → {len(filtradas)} linhas aceitas")
        return filtradas

    except HTTPException:
        # deixa FastAPI propagar como está
        raise
    except Exception as ex:
        print("[dados_core] ERRO BRUTO:", repr(ex))
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=(
                f"Falha interna no filtro de linhas: {type(ex).__name__}: {ex}"
            ),
        )


# ---------------------------------------------------------------------
# /api/dados/linhas  (compatível com chamadas antigas)
# ---------------------------------------------------------------------

@app.get("/api/dados/linhas")
def dados_linhas(
    etapa: str = Query(..., description="Etapa (e.g., fundamental_I)"),
    disciplina: str = Query(..., description="Nome da disciplina"),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    conteudo: Optional[str] = Query(None),
    aula: Optional[str] = Query(None),
    habilidade: Optional[str] = Query(
        None,
        description="Filtro opcional por habilidade BNCC (códigos separados por ||)."
    ),
    contains: bool = Query(False, description="0 (AND) ou 1 (OR) para filtros múltiplos"),
):
    """
    Retorna LINHAS (registros crus) filtradas por contexto.
    Mantido para compatibilidade com chamadas antigas.
    """
    return _filtrar_linhas_por_contexto(
        etapa=etapa,
        disciplina=disciplina,
        tema=tema,
        objeto=objeto,
        titulo=titulo,
        conteudo=conteudo,
        aula=aula,
        habilidade=habilidade,
        contains=contains,
    )


# ---------------------------------------------------------------------
# /api/dados/conteudos  (usado pelo TextareaUI para montar o plano)
# ---------------------------------------------------------------------

@app.get("/api/dados/conteudos")
def dados_conteudos(
    etapa: str = Query(..., description="Etapa (e.g., fundamental_I)"),
    disciplina: str = Query(..., description="Nome da disciplina"),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    conteudo: Optional[str] = Query(None),
    aula: Optional[str] = Query(None),
    habilidade: Optional[str] = Query(
        None,
        description="Filtro opcional por habilidade BNCC (códigos separados por ||)."
    ),
    contains: bool = Query(False, description="0 (AND) ou 1 (OR) para filtros múltiplos"),
):
    """
    Retorna as mesmas LINHAS cruas que /api/dados/linhas,
    mas é o endpoint pensado para:
      - montar select de conteúdos
      - alimentar textareas (Objetivos / Conhecimentos Prévios / etc.)
    """
    return _filtrar_linhas_por_contexto(
        etapa=etapa,
        disciplina=disciplina,
        tema=tema,
        objeto=objeto,
        titulo=titulo,
        conteudo=conteudo,
        aula=aula,
        habilidade=habilidade,
        contains=contains,
    )





@app.get("/api/auto/sugerir")
def auto_sugerir(
    etapa: str = Query(...),
    disciplina: str = Query(...),
    q: str = Query(..., description="termo: parte do tema, palavra-chave, ou código de habilidade"),
):
    """
    Sugere (tema, partes[], habilidades[]) a partir de 'q'.
    Estratégia: busca 'q' (case-insensitive) em Tema, Objeto do Conhecimento, Habilidades e Títulos.
    O tema mais recorrente entre os matches vira a sugestão principal.
    """
    arquivo = _find_arquivo(etapa, disciplina)
    data = _safe_json(_safe_path(arquivo)) or []

    qn = _normalize(q)

    def _val(row, keys):
        for k in keys:
            if k in row and str(row[k]).strip():
                return str(row[k]).strip()
        return None

    # coleta matches por tema
    temas_hit: Dict[str, Dict[str, Any]] = {}  # tema -> { partes:set, habilidades:set, score:int }
   # Se já existir TITLE_KEYS global, reaproveite:
    try:
        TITULO_KEYS = TITLE_KEYS
    except NameError:
        TITULO_KEYS = ["Título da aula","TÍTULO DA AULA","Titulo da aula","Título","TÍTULO","Titulo"]

    # ...
    objeto = _val(row, OBJ_KEYS)  # usar chaves de Objeto do Conhecimento



    def _hit_in(s: Optional[str]) -> bool:
        if not s: return False
        return qn in _normalize(s)

    for row in data:
        tema_row = _val(row, TEMA_KEYS)
        if not tema_row:
            continue
        objeto = _val(row, CONTEUDO_KEYS)
        hab    = _val(row, HAB_KEYS)
        titulo = None
        for tk in TITULO_KEYS:
            if tk in row and str(row[tk]).strip():
                titulo = str(row[tk]).strip()
                break

        # pontua se 'q' aparecer em qualquer campo relevante
        score = 0
        for cand in [tema_row, objeto, hab, titulo]:
            if _hit_in(cand):
                score += 1

        if score == 0:
            continue

        rec = temas_hit.setdefault(tema_row, {"partes": set(), "habilidades": set(), "score": 0})
        rec["score"] += score
        if titulo:
            rec["partes"].add(titulo)
        if hab:
            rec["habilidades"].add(hab)

    if not temas_hit:
        return {"tema": None, "partes": [], "habilidades": []}

    # escolhe o tema com maior 'score'
    tema_sug = max(temas_hit.items(), key=lambda kv: kv[1]["score"])[0]
    partes_list = sorted(temas_hit[tema_sug]["partes"]) if temas_hit[tema_sug]["partes"] else []
    habs_list   = sorted(temas_hit[tema_sug]["habilidades"]) if temas_hit[tema_sug]["habilidades"] else []

    # como bônus, se não achou partes via 'q', tenta derivar do endpoint de partes:
    if not partes_list:
        try:
            # Reaproveita a lógica do próprios dados (sem fazer HTTP)
            partes: Dict[str, Dict[str, Any]] = {}
            tema_norm = re.sub(r"\s+", " ", tema_sug.strip().lower())
            for row in data:
                row_tema = _val(row, TEMA_KEYS)
                if not row_tema or re.sub(r"\s+", " ", row_tema.strip().lower()) != tema_norm:
                    continue
                titulo = None
                for tk in TITULO_KEYS:
                    if tk in row and str(row[tk]).strip():
                        titulo = str(row[tk]).strip(); break
                if titulo:
                    partes.setdefault(titulo, True)
            if partes:
                partes_list = sorted(partes.keys())
        except Exception:
            pass

    return {"tema": tema_sug, "partes": partes_list, "habilidades": habs_list}

@app.get("/api/dados/partes")
def dados_partes(
    arquivo: Optional[str] = Query(None),
    etapa:   Optional[str] = Query(None),
    disciplina: Optional[str] = Query(None),
    tema: Optional[str] = Query(None),
    objeto: Optional[str] = Query(None),
    titulo: Optional[str] = Query(None),
    contains: Optional[str] = Query(None),
):
    """
    Retorna “partes” da aula para o Tema/Objeto/Título selecionados.
    Implementação: usa os mesmos títulos do dataset e consolida por:
      - “Prefixo – Partes 1, 2 …” quando os títulos seguem o padrão “... – Parte N”
      - Caso não haja “Parte N”, devolve títulos normais (sem agrupar), 1 por linha.
    """
    try:
        e_norm = normalize_etapa(etapa)
        d_norm = normalize_disciplina_alias(e_norm, (disciplina or "")) if disciplina else None
        contains_b = _flex_bool(contains)

        # sem disciplina e sem arquivo → nada
        if not d_norm and not arquivo:
            return {"partes": []}

        # filtros (placeholders já são descartados por _normset_pipe)
        sel_tema     = _normset_pipe(tema)
        sel_obj      = _normset_pipe(objeto)
        sel_titulo   = _normset_pipe(titulo)

        rows = _load_rows_canon(e_norm, d_norm, arquivo)

        # Se não tem filtro nenhum, devolve consolidação por títulos (todas as partes)
        has_extra = any([sel_tema, sel_obj, sel_titulo])
        if not has_extra:
            # usa o mesmo agrupador da prévia
            saida = _compose_agrupar_titulos(rows, FIELD_MAP, e_norm, d_norm)
            return {"partes": finalize_list(saida)}

        # Com filtros finos:
        grupo: list[dict] = []
        for r in rows:
            try:
                row_etapa = str(_get_field_ci(r, ALIASES["etapa"]) or "").strip()
                if row_etapa and norm_key_ci(row_etapa) != norm_key_ci(e_norm or row_etapa):
                    continue
            except Exception:
                pass

            temas    = row_temas(r, FIELD_MAP, e_norm, d_norm)
            objetos  = row_objetos(r, FIELD_MAP, e_norm, d_norm)
            titulos  = row_titles(r, FIELD_MAP, e_norm, d_norm)

            if sel_tema   and not _match_list(temas,   sel_tema,   contains_b): continue
            if sel_obj    and not _match_list(objetos, sel_obj,    contains_b): continue
            if sel_titulo and not _match_list(titulos, sel_titulo, contains_b): continue

            grupo.append(r)

        # Consolida em “Partes …” quando aplicável
        saida = _compose_agrupar_titulos(grupo, FIELD_MAP, e_norm, d_norm)
        return {"partes": finalize_list(saida)}
    except Exception as e:
        print("[/api/dados/partes] ERRO:", repr(e))
        raise HTTPException(status_code=422, detail=f"Falha ao processar /api/dados/partes: {e.__class__.__name__}: {e}")


def _extract_codigo_from_habilidade(hab_txt: str) -> Optional[str]:
    if not hab_txt:
        return None
    m = CODIGO_RE.match(hab_txt.strip())
    if not m:
        # tenta pegar um código em qualquer lugar (mais “solto”)
        m = re.search(r'\b[A-Z]{2,}\d{2,}[A-Z]*\d*\b', hab_txt)
    return (m.group(0).strip(" :-()") if m else None)


def _payload_filters(payload: Dict[str, Any]) -> dict:
    """Lê os filtros enviados pelo front (se existirem)."""
    def _normset(val):
        arr = []
        if isinstance(val, str):
            arr = [x.strip() for x in val.split("||") if x.strip()]
        elif isinstance(val, list):
            arr = [str(x).strip() for x in val if str(x).strip()]
        return {_norm_key(x) for x in arr}

    # aceitamos várias chaves: “tema” ou “temas”, “titulo(s)”, “habilidade(s)”, “objetivos_especificos”
    temas    = _normset(payload.get("tema") or payload.get("temas"))
    titulos  = _normset(payload.get("titulos") or payload.get("titulos_da_aula") or _titles_from_payload(payload))
    habs_sel = {s.strip() for s in _flatten_lines(payload.get("habilidades") or payload.get("habilidade") or [])}
    oes_sel  = _normset(payload.get("objetivos_especificos") or payload.get("objetivo_especifico"))

    return {"temas": temas, "titulos": titulos, "habs": habs_sel, "oes": oes_sel}

def _group_oa_by_hab_using_json(payload: Dict[str, Any]) -> Dict[str, List[str]]:
    """
    Gera: { "<codigo ou label da habilidade>": ["OA normalizado (única linha)"], ... }
    Regra:
      - Agrupa por HABILIDADE.
      - Usa exclusivamente o campo 'objetivos' do JSON (normalização leve).
      - Se NÃO houver 'objetivos' em um grupo, aplica Fallback:
            "Habilidade — {objetos do conhecimento}; {conteúdos} — base: {títulos}"
        (omitindo segmentos inexistentes no JSON).
    """
    etapa   = (payload.get("etapa") or "").strip() or None
    disc    = (payload.get("componente_curricular") or payload.get("disciplina") or "").strip() or None
    arquivo = payload.get("arquivo") or None

    data = _load_json_by(etapa=etapa, disciplina=disc, arquivo=arquivo) or []
    data = _canonize_data_rows(data)

    f = _payload_filters(payload)
    temas_sel   = f["temas"]
    titulos_sel = f["titulos"]
    habs_sel    = f["habs"]  # códigos (ex.: EF08MA08)

    # mapa: codigo_ou_label -> {"hab_txt": str, "oa": list[str], "rows": list[dict]}
    bucket: Dict[str, Dict[str, Any]] = {}

    for row in data if isinstance(data, list) else []:
        # Tema (precisa de FIELD_MAP/etapa/disc)
        row_temas_norm = {_norm_key(x) for x in row_temas(row, FIELD_MAP, etapa or "", disc or "")}
        if temas_sel and not row_temas_norm.intersection(temas_sel):
            continue


        # Título (idem)
        if titulos_sel:
            row_titles_norm = {_norm_key(t) for t in _row_titles(row, FIELD_MAP, etapa, disc)}
            if not row_titles_norm.intersection(titulos_sel):
                continue

        # Habilidade (texto bruto) + código
        hab_txt = _get_field_ci(row, HAB_KEYS) or ""
        if not hab_txt.strip():
            continue
        code = _extract_codigo_from_habilidade(hab_txt)
        key  = code or hab_txt.strip()
        if habs_sel and code and code.upper() not in habs_sel:
            continue

        # Objetivos (normalizados) desta linha
        oa_list = _normalize_objetivos(
            _get_field_ci(row, OBJ_APR_PRIMARY) or _get_field_ci(row, OBJ_APR_FALLBACK) or ""
        )

        # Carregar registro do bucket
        rec = bucket.setdefault(key, {"hab_txt": hab_txt.strip(), "oa": [], "rows": []})
        if oa_list:
            rec["oa"].extend(oa_list)
        rec["rows"].append(row)

    # Consolidação final: 1 linha de OA por habilidade
    out: Dict[str, List[str]] = {}
    for key, rec in bucket.items():
        oa_norm = _unique_preserve_order(rec["oa"])
        if oa_norm:
            out[key] = [" — ".join(oa_norm)]
        else:
            # Fallback: montar linha com objetos/conteúdos/títulos (somente se 'objetivos' ausentes)
            objs = _unique_preserve_order(
                sum((_row_objetos(r, FIELD_MAP, etapa, disc) for r in rec["rows"]), [])
            )
            cts = _unique_preserve_order(
                sum((_row_conteudos(r, FIELD_MAP, etapa, disc) for r in rec["rows"]), [])
            )
            # Chame com field_map/etapa/disc conforme a nova assinatura
            tits = _unique_preserve_order(
                _compose_agrupar_titulos(rec["rows"], FIELD_MAP, etapa, disc)
                if _compose_agrupar_titulos.__code__.co_argcount >= 4
                else _compose_agrupar_titulos(rec["rows"])
            )

            parts = []
            # Habilidade já é o título da linha; partes complementares:
            if objs:
                parts.append("; ".join(objs))
            if cts:
                parts.append("; ".join(cts))
            if tits:
                parts.append("base: " + ", ".join(tits))
            out[key] = [" — ".join([rec["hab_txt"]] + parts) if parts else rec["hab_txt"]]

    return out



# === COMPOSE: prévia consolidada (habilidades + OA + títulos) ===


# Detecta "… – Parte X"
_COMPOSE_RE_PARTE = re.compile(r"^(.*?)(?:\s*[-–]\s*Parte\s*(\d+))\s*$", re.IGNORECASE)

def _compose_int_aula(val) -> Optional[int]:
    """Extrai número da aula (int) de campos variados."""
    if val is None:
        return None
    s = str(val).strip()
    # tenta primeiro dígito(s) (ex.: "3", "03", "3A", "3 - dupla")
    m = re.match(r"^\s*(\d+)", s)
    if m:
        try:
            return int(m.group(1))
        except Exception:
            return None
    try:
        return int(float(s.replace(",", ".").split()[0]))
    except Exception:
        return None

def _compose_get_title(row: dict) -> str:
    """Obtém o Título cru (sem anexar ', Aula N'). Reaproveita suas chaves conhecidas."""
    for k in ["Título da aula","TÍTULO DA AULA","Titulo da aula","Título","TÍTULO","Titulo","TÍTULO "]:
        if k in row and str(row[k]).strip():
            return str(row[k]).strip()
    # fallback: se não houver título, tenta usar o Tema
    val = _get_field_ci(row, TEMA_KEYS)
    return val or ""

def _compose_format_list_pt(nums: list[int], rotulo: str) -> str:
    """Formata 'Aulas 1, 2 e 3'."""
    ns = sorted({int(n) for n in nums if isinstance(n, (int, str)) and str(n).isdigit()})
    if not ns:
        return ""
    if len(ns) == 1:
        return f"{rotulo} {ns[0]}"
    frente = ", ".join(map(str, ns[:-1]))
    return f"{rotulo} {frente} e {ns[-1]}"

def _compose_agrupar_titulos(rows: list[dict], field_map: dict, etapa: str | None = None, disc: str | None = None) -> list[str]:
    """
    Consolida títulos respeitando aliases do field_map:
      - Se casar '... – Parte X' → 'prefixo – Partes 1, 2 e 3'
      - Caso contrário → 'Título – Aulas 4 e 7'
    Obs.: para grupos com 'Partes', não adicionamos 'Aulas ...'.
    """
    grupos_parte = defaultdict(lambda: {"prefixo": None, "partes": set()})
    grupos_titulo = defaultdict(lambda: {"titulo": None, "aulas": set()})

    for r in (rows or []):
        # 1) TÍTULO a partir do extrator oficial
        titles = _row_titles(r, field_map, etapa, disc) or []
        titulo = next((t for t in titles if str(t).strip()), "").strip()

        # fallback (mantém compat): se por algum motivo vier vazio, usa helper antigo
        if not titulo and '_compose_get_title' in globals():
            try:
                titulo = (_compose_get_title(r) or "").strip()
            except Exception:
                pass

        if not titulo:
            continue

        # 2) AULA a partir do extrator oficial; fallback para colunas cruas
        aulas = _row_aulas(r, field_map, etapa, disc) or []
        if not aulas:
            aulas = [r.get("AULA"), r.get("Aula"), r.get("n_aula")]
            aulas = [x for x in aulas if x is not None]

        # normaliza cada possível aula para inteiro
        aula_nums = set()
        for a in aulas:
            n = _compose_int_aula(a)
            if n is not None:
                aula_nums.add(n)

        # 3) Agrupamento por "Parte X"
        m = _COMPOSE_RE_PARTE.match(titulo)
        if m:
            pref, num = m.group(1).strip(), m.group(2)
            g = grupos_parte[pref]
            g["prefixo"] = pref
            if num and str(num).isdigit():
                g["partes"].add(int(num))
        else:
            g = grupos_titulo[titulo]
            g["titulo"] = titulo
            for n in aula_nums:
                g["aulas"].add(n)

    saida = []

    # 1) grupos com "Parte X"
    for g in grupos_parte.values():
        partes = sorted(g["partes"])
        if len(partes) > 1:
            titulo_fmt = f"{g['prefixo']} – Partes {', '.join(map(str, partes[:-1]))} e {partes[-1]}"
        elif len(partes) == 1:
            titulo_fmt = f"{g['prefixo']} – Parte {partes[0]}"
        else:
            titulo_fmt = g["prefixo"]
        saida.append(titulo_fmt)

    # 2) grupos por TÍTULO (sem "Parte X")
    for g in grupos_titulo.values():
        t = g["titulo"]
        aulas_txt = _compose_format_list_pt(sorted(g["aulas"]), "Aulas")
        saida.append(f"{t} – {aulas_txt}" if aulas_txt else t)

    # dedupe estável (preserva ordem)
    dedup, seen = [], set()
    for s in saida:
        if s not in seen:
            seen.add(s); dedup.append(s)
    return dedup


@app.post("/api/compose/plano")
def compose_plano(payload: dict = Body(...)):
    """
    Corpo esperado:
      {
        "aulas": [ { ...linha bruta do JSON da disciplina... }, ... ]
      }
    Retorna:
      {
        "habilidades": [texto completo...],
        "oa": ["OA (unidos) por habilidade...", ...],
        "titulos": ["Prefixo – Partes ...", "Título – Aulas ...", ...],
        "detalhe_aulas": [{aula, titulo, objeto, conteudo, habilidade, codigo?}, ...]
      }
    """
    aulas = payload.get("aulas") if isinstance(payload, dict) else None
    if not isinstance(aulas, list) or not aulas:
        return {"habilidades": [], "oa": [], "titulos": [], "detalhe_aulas": []}

    # ---- Agrupa por habilidade (texto completo)
    by_hab = defaultdict(lambda: {"hab_txt": "", "oa": set(), "rows": []})
    detalhe = []

    for row in aulas:
        if not isinstance(row, dict):
            continue

        hab_txt = _get_field_ci(row, HAB_KEYS) or ""   # texto integral
        if not hab_txt.strip():
            continue

        # OA desta linha (pode vir multi-linha; helper já quebra e limpa)
        oa_list = _row_objetivos_aprendizagem(row)

        # Armazena
        d = by_hab[hab_txt.strip()]
        d["hab_txt"] = hab_txt.strip()
        for oa in oa_list:
            if oa.strip():
                d["oa"].add(oa.strip())
        d["rows"].append(row)

        # Detalhamento da aula para a prévia
        detalhe.append({
            "aula": _compose_int_aula(row.get("AULA") or row.get("Aula") or row.get("n_aula")),
            "titulo": _compose_get_title(row),
            "objeto": _get_field_ci(row, OBJ_KEYS),
            "conteudo": _get_field_ci(row, CONTEUDO_KEYS),  # usar chaves canônicas
            "habilidade": hab_txt.strip()
        })


    # ---- Monta listas finais
    habilidades_out: list[str] = []
    oa_out: list[str] = []
    titulos_out: list[str] = []

    for hab_key, bloc in by_hab.items():
        if not bloc["hab_txt"]:
            continue
        habilidades_out.append(bloc["hab_txt"])  # texto completo
        # OA em UMA linha por habilidade
        oa_out.append(" — ".join(sorted(bloc["oa"], key=str.lower)))
        # Títulos consolidados (linhas que pertencem a esta habilidade)
        titulos_out.extend(_compose_agrupar_titulos(bloc["rows"], FIELD_MAP))
    # se tiver etapa/disc acessíveis aqui, pode usar:
    # titulos_out.extend(_compose_agrupar_titulos(bloc["rows"], FIELD_MAP, etapa, disc))


    # dedupe suave preservando ordem
    def _dedup(seq):
        seen = set(); out = []
        for x in seq:
            if x not in seen:
                seen.add(x); out.append(x)
        return out

    return {
        "habilidades": _dedup(habilidades_out),
        "oa": _dedup(oa_out),
        "titulos": _dedup(titulos_out),
        "detalhe_aulas": sorted(
            detalhe,
            key=lambda r: (999999 if r["aula"] is None else r["aula"], _norm_key(r.get("titulo") or ""))
        )
    }

def _ensure_preview_parity(payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normaliza o payload para que Export == Prévia:
      - titulos_da_aula: enriquecidos com ", Aula N"
      - objetivos_aprendizagem: se vazio, gera via _group_oa_by_hab_using_json (1 linha por habilidade)
      - conteudos_habilidades_bncc: se vazio, monta a partir de 'linhas' (código + descrição)
    Não altera campos já preenchidos.
    """
    p = payload or {}

    # 1) Títulos
    try:
        tit_norm = _titles_for_export(p)  # já garante ", Aula N" quando possível
    except Exception:
        tit_norm = _flatten_lines(p.get("titulos_da_aula"))
    if tit_norm and not _flatten_lines(p.get("titulos_da_aula")):
        p["titulos_da_aula"] = tit_norm

    # 2) OA (uma linha por habilidade)
    oa_list = _flatten_lines(p.get("objetivos_aprendizagem"))
    if not oa_list:
        try:
            oa_grouped = _group_oa_by_hab_using_json(p)  # {code: [linha1,...]}
            oa_lines = []
            for code, linhas in oa_grouped.items():
                linha = " — ".join(_flatten_lines(linhas))
                if linha:
                    oa_lines.append(linha)
            p["objetivos_aprendizagem"] = oa_lines
        except Exception:
            # fallback: deixa vazio se não conseguir inferir
            p.setdefault("objetivos_aprendizagem", [])

    # 3) Habilidades (BNCC) — texto integral
    bncc_txt = p.get("conteudos_habilidades_bncc") or ""
    if not bncc_txt:
        linhas = p.get("linhas") or []
        linhas_txt = []
        for ln in linhas:
            if not isinstance(ln, dict): 
                continue
            codigo = (ln.get("codigo") or "").strip()
            habtx  = (ln.get("habilidade") or "").strip()
            if codigo or habtx:
                linhas_txt.append(f"{codigo} {habtx}".strip())
        if linhas_txt:
            p["conteudos_habilidades_bncc"] = "\n".join(_unique_preserve_order(linhas_txt))

    return p


# ---------------------------------------------------------------------
# Rotas – Exportações (JSON, DOCX, PPTX, XLSX, PDF)
# Todas aceitam o MESMO corpo JSON (payload do plano)
# ---------------------------------------------------------------------
# === COLE ESTE BLOCO NOVO (rotas de exportação) ===

def _titles_for_export(payload: Dict[str, Any]) -> list[str]:
    """
    Devolve os títulos padronizados para exportação.
    Regras:
      1) Se o payload trouxer títulos, enriquecemos com ", Aula N" quando faltar.
      2) Se o payload não trouxer títulos, carregamos do JSON (filtrado por temas)
         — nesses casos, _row_titles já vem no formato "Título, Aula N".
    """
    # --------- Helpers locais ---------
    def _int_aula(val) -> Optional[int]:
        if val is None:
            return None
        s = str(val).strip()
        m = re.match(r"^\s*(\d+)", s)
        if m:
            try:
                return int(m.group(1))
            except Exception:
                return None
        try:
            return int(float(s.replace(",", ".").split()[0]))
        except Exception:
            return None

    def _raw_title_from_row(row: dict) -> str:
        # procura título em várias chaves conhecidas (sem anexar ", Aula N")
        for k in ["Título da aula","TÍTULO DA AULA","Titulo da aula","Título","TÍTULO","Titulo","TÍTULO "]:
            if k in row and str(row[k]).strip():
                return str(row[k]).strip()
        # fallback por tema (melhor do que nada)
        return _get_field_ci(row, TEMA_KEYS) or ""

    # detecta qualquer variante de "Aula N" (vírgula, hífen, dois-pontos, espaços)
    aula_any_re = re.compile(r"(?:,|–|-|:)?\s*Aula\s*(\d+)\s*$", re.I)

    # --------- Dados de contexto ---------
    etapa = (payload.get("etapa") or "").strip() or None
    disc  = (payload.get("componente_curricular") or payload.get("disciplina") or "").strip() or None
    arquivo = payload.get("arquivo") or None

    # 1) Títulos vindos do payload (como o front enviou)
    tit_payload = _flatten_lines(_titles_from_payload(payload))

    # 2) Se não há nada no payload → comportamento anterior (carrega do JSON)
    if not tit_payload:
        try:
            data = _load_json_by(etapa=etapa, disciplina=disc, arquivo=arquivo) or []
            data = _canonize_data_rows(data)
        except Exception:
            data = []
        temas_sel = {_norm_key(t) for t in _flatten_lines(payload.get("tema") or payload.get("temas") or [])}
        out: list[str] = []
        seen_norm: set[str] = set()
        for row in data if isinstance(data, list) else []:
            row_temas_norm = {_norm_key(x) for x in row_temas(row, FIELD_MAP, etapa or "", disc or "")}
            if temas_sel and not row_temas.intersection(temas_sel):
                continue
            for t in _row_titles(row, FIELD_MAP, etapa or "", disc or ""):  # já "Título, Aula N"
                key_norm = _norm_key(t)
                if key_norm not in seen_norm:
                    seen_norm.add(key_norm)
                    out.append(t)

        return out

    # 3) Há títulos no payload → enriquecer (anexar ", Aula N" quando faltar)
    #    Para isso montamos um índice a partir do JSON da disciplina.
    try:
        data = _load_json_by(etapa=etapa, disciplina=disc, arquivo=arquivo) or []
        data = _canonize_data_rows(data)
    except Exception:
        data = []

    # índice: título cru normalizado -> {aulas: set(int), title: grafia de referência}
    idx: Dict[str, Dict[str, Any]] = {}
    if isinstance(data, list):
        for row in data:
            raw_title = _raw_title_from_row(row)
            if not raw_title:
                continue
            n_aula = _int_aula(row.get("AULA") or row.get("Aula") or row.get("n_aula"))
            key = _norm_key(raw_title)
            rec = idx.setdefault(key, {"title": raw_title, "aulas": set()})
            if n_aula is not None:
                rec["aulas"].add(n_aula)

    enriched: list[str] = []
    seen_norm: set[str] = set()

    for t in tit_payload:
        base_in = str(t or "").strip()
        if not base_in:
            continue

        # Tenta extrair "Aula N" caso já venha consigo
        m = aula_any_re.search(base_in)
        aula_from_payload = int(m.group(1)) if m else None

        # Remove o sufixo "Aula N" para descobrir o "título base" que usaremos no match
        base_title = aula_any_re.sub("", base_in).rstrip(", ").rstrip("–").rstrip("-").rstrip(":").strip()

        key = _norm_key(base_title) if base_title else _norm_key(base_in)
        rec = idx.get(key)

        if rec and rec["aulas"]:
            # Temos referência de aulas no JSON — Regra 1: completar as que faltam
            base_ref = rec["title"]  # usa grafia do JSON
            aulas_sorted = sorted(rec["aulas"])
            if aula_from_payload is not None and aula_from_payload not in aulas_sorted:
                # se a aula do payload não consta do índice, ainda assim preservamos essa entrada específica
                line = f"{base_ref}, Aula {aula_from_payload}"
                kn = _norm_key(line)
                if kn not in seen_norm:
                    seen_norm.add(kn)
                    enriched.append(line)
            # Adiciona todas as aulas conhecidas (completa o que estiver faltando)
            for n in aulas_sorted:
                line = f"{base_ref}, Aula {n}"
                kn = _norm_key(line)
                if kn not in seen_norm:
                    seen_norm.add(kn)
                    enriched.append(line)
        else:
            # Sem correspondência no índice — preserva como veio
            # (se tiver "Aula N", mantém; senão mantém cru)
            keep = base_in
            kn = _norm_key(keep)
            if kn not in seen_norm:
                seen_norm.add(kn)
                enriched.append(keep)

    return enriched



def _safe_payload(payload: Dict[str, Any]) -> Dict[str, Any]:
    # Garante chaves básicas esperadas
    payload = payload or {}
    payload.setdefault("identificacao", "Plano de Aula")
    payload.setdefault("disciplina", payload.get("componente_curricular", ""))
    payload.setdefault("temas", payload.get("tema", []))
    payload.setdefault("linhas", payload.get("linhas", []))
    return payload

@app.post("/export/docx")
def export_docx(payload: Dict[str, Any] = Body(...)):
    # --- Dependências obrigatórias ---
    try:
        import io
        from docx import Document
        from docx.shared import Pt
        from fastapi import HTTPException
        from fastapi.responses import StreamingResponse
    except Exception as e:
        # Se python-docx não existir, cai no except abaixo
        pass

    try:
        from docx import Document  # tenta de novo para mensagem clara
        from docx.shared import Pt
    except Exception:
        raise HTTPException(
            status_code=501,
            detail="python-docx não instalado. Instale com: pip install python-docx"
        )

    # === [PATCH] aceitar novo payload sem quebrar o legado ===
    parsed = None
    if isinstance(payload, dict):
        try:
            parsed  = ExportPayload(**payload)
            adapted = _to_legacy_dict(parsed)
            payload = {**payload, **adapted}
        except Exception:
            parsed = None

    # Reflete "aulas" do modelo novo, se houver
    if parsed and getattr(parsed, "aulas", None):
        payload = {**payload, "aulas": parsed.aulas}

    # ---------- Helpers locais ----------
    def _flat_lines(x):
        if x is None:
            return []
        if isinstance(x, (list, tuple, set)):
            out = []
            for it in x:
                s = str(it).strip()
                if not s:
                    continue
                out.extend([t for t in s.split("\n") if t.strip()])
            return out
        return [t for t in str(x).split("\n") if t.strip()]

    def _flat_first_nonempty(*vals):
        """Retorna a 1ª lista 'flattened' não vazia dentre várias chaves/aliases."""
        for v in vals:
            fl = _flat_lines(v)
            if fl:
                return fl
        return []

    def _flat_text(x) -> str:
        """Transforma lista em texto com quebras; string vira strip."""
        if isinstance(x, (list, tuple, set)):
            return "\n".join([s for s in _flat_lines(x)])
        return "\n".join(_flat_lines(x))

    # ---------- Normalização do payload ----------
    p = payload or {}
    p = _ensure_preview_parity(p)

    ident = (p.get("identificacao") or "Plano de Aula")

    etapa = (p.get("etapa") or "").strip()
    disc  = (p.get("componente_curricular") or p.get("disciplina") or "").strip()

    # temas pode vir como 'tema' (str/list), 'temas' (str/list)
    temas = _flat_first_nonempty(p.get("tema"), p.get("temas"))

    # títulos já enriquecidos "Título, Aula N"
    try:
        titulos = _titles_for_export(p)
    except Exception:
        titulos = []

    # conteudos pode vir como 'conteudo'/'conteudos'
    conteudos = _flat_first_nonempty(p.get("conteudos"), p.get("conteudo"))

    # Seção habilidades: usa textarea se vier; senão monta de linhas (código + descrição)
    # 1) Habilidades: prioridade para o textarea de habilidades
    bncc_txt = _flat_text(
        p.get("habilidades_texto")
        or p.get("habilidades_bncc")
    )

    # 2) Se não tiver, usa o campo combinado (compat)
    if not bncc_txt:
        bncc_txt = _flat_text(p.get("conteudos_habilidades_bncc") or "")

    # 3) Se ainda assim estiver vazio, fallback nas linhas cruas do JSON
    if not bncc_txt:
        linhas = p.get("linhas") or []
        linhas_txt = []
        for ln in linhas:
            try:
                codigo = (ln.get("codigo") or "").strip()
                habtx  = (ln.get("habilidade") or "").strip()
                if codigo or habtx:
                    linhas_txt.append(f"{codigo} {habtx}".strip())
            except Exception:
                continue
        bncc_txt = "\n".join(_unique_preserve_order(linhas_txt))


    obj_espec  = p.get("objetivos_especificos") or ""
    previos    = p.get("conhecimentos_previos") or ""
    metod      = p.get("metodologia_estrategias") or []
    recs       = p.get("recursos_didaticos") or []
    atividades = p.get("atividades_desempenho") or []
    criterios  = p.get("criterios_avaliacao") or []
    autoav     = p.get("avaliacao_autoavaliacao") or []
    consid     = p.get("consideracoes_finais") or []

    # >>> Objetivos de Aprendizagem GERADOS do JSON conforme filtros (defensivo)
    try:
        oa_grouped = _group_oa_by_hab_using_json(p) or {}  # {code: [linha1, linha2]}
    except Exception:
        oa_grouped = {}

    # ---------- Início DOCX ----------
    buf = io.BytesIO()
    doc = Document()

    # ---- Arial 12 como padrão
    st = doc.styles['Normal']
    st.font.name = 'Arial'
    st.font.size = Pt(12)

    # Estilo dos títulos
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(12)
    style_h1.font.bold = True

    # Título principal
    style_title = doc.styles['Title']
    style_title.font.name = 'Arial'
    style_title.font.size = Pt(14)
    style_title.font.bold = True

    doc.add_heading("Plano de Aula", level=0)

    # 1) Identificação
    doc.add_heading("1. Identificação", level=1)
    for line in _flat_lines(ident):
        doc.add_paragraph(line)

    # 2) Contexto
    doc.add_heading("2. Contexto", level=1)
    if etapa:
        doc.add_paragraph(f"Etapa: {etapa}")
    if disc:
        doc.add_paragraph(f"Componente Curricular: {disc}")
    if temas:
        doc.add_paragraph("Tema(s): " + ", ".join(_unique_preserve_order([t.strip() for t in temas])))

    if titulos:
        doc.add_paragraph("Título(s) da Aula:")
        for t in _unique_preserve_order([str(x).strip() for x in titulos if str(x).strip()]):
            doc.add_paragraph(t, style="List Bullet")

    # 3) Objetivos Específicos
    oe = _flat_lines(obj_espec)
    if oe:
        doc.add_heading("3. Objetivos Específicos", level=1)
        for v in _unique_preserve_order(oe):
            doc.add_paragraph(v, style="List Bullet")

    # 4) Objetivos de Aprendizagem (agrupado)
    if oa_grouped:
        doc.add_heading("4. Objetivos de Aprendizagem", level=1)
        # ordena por código para previsibilidade
        for code in sorted(oa_grouped.keys()):
            linhas = [s.strip() for s in _flat_lines(oa_grouped.get(code))]
            if not linhas:
                continue
            par = doc.add_paragraph()
            run = par.add_run(f"{code}: ")
            run.bold = True
            # separa por "; " para reduzir risco de linhas longas com travessão
            par.add_run("; ".join(_unique_preserve_order(linhas)))

    # 5) Conteúdos e Habilidades (BNCC)
    cont = _flat_lines(conteudos)
    bncc = _flat_lines(bncc_txt)
    if cont or bncc:
        doc.add_heading("5. Conteúdos e Habilidades (BNCC)", level=1)
        if cont:
            doc.add_paragraph("Conteúdos:")
            for v in _unique_preserve_order(cont):
                doc.add_paragraph(v, style="List Bullet")
        if bncc:
            doc.add_paragraph("Habilidades (BNCC):")
            for v in _unique_preserve_order(bncc):
                doc.add_paragraph(v)

    # 6) Conhecimentos Prévios
    cp = _flat_lines(previos)
    if cp:
        doc.add_heading("6. Conhecimentos Prévios", level=1)
        for v in _unique_preserve_order(cp):
            doc.add_paragraph(v)

    # 7) Metodologia e Estratégias
    met = _flat_lines(metod)
    if met:
        doc.add_heading("7. Metodologia e Estratégias", level=1)
        for v in _unique_preserve_order(met):
            doc.add_paragraph(v, style="List Bullet")

    # 8) Recursos Didáticos
    recs = _flat_lines(recs)
    if recs:
        doc.add_heading("8. Recursos Didáticos", level=1)
        for v in _unique_preserve_order(recs):
            doc.add_paragraph(v, style="List Bullet")

    # 9) Atividades de Desempenho
    atividades = _flat_lines(atividades)
    if atividades:
        doc.add_heading("9. Atividades de Desempenho", level=1)
        for v in _unique_preserve_order(atividades):
            doc.add_paragraph(v, style="List Bullet")

    # 10) Critérios de Avaliação
    criterios = _flat_lines(criterios)
    if criterios:
        doc.add_heading("10. Critérios de Avaliação", level=1)
        for v in _unique_preserve_order(criterios):
            doc.add_paragraph(v, style="List Bullet")

    # 11) Avaliação e Autoavaliação
    autoav = _flat_lines(autoav)
    if autoav:
        doc.add_heading("11. Avaliação e Autoavaliação", level=1)
        for v in _unique_preserve_order(autoav):
            doc.add_paragraph(v, style="List Bullet")

    # 12) Considerações Finais
    cf = _flat_lines(consid)
    if cf:
        doc.add_heading("12. Considerações Finais", level=1)
        for v in _unique_preserve_order(cf):
            doc.add_paragraph(v)

    # ---------- Finaliza ----------
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{_slug(disc or 'Plano')}.docx"
    headers = {"Content-Disposition": f'attachment; filename=\"{filename}\"'}
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers
    )

@app.post("/export/json")
def export_json(payload: Dict[str, Any] = Body(...)):
    # 1) Compat: aceitar novo formato e adaptar sem quebrar legado
    if isinstance(payload, dict):
        try:
            parsed  = ExportPayload(**payload)
            adapted = _to_legacy_dict(parsed)
            payload = {**payload, **adapted}
            # refletir "aulas" do modelo novo, se existir
            if getattr(parsed, "aulas", None):
                payload["aulas"] = parsed.aulas
        except Exception:
            pass

    # 2) Paridade com prévia (mesmas regras do DOCX/PDF)
    p = _ensure_preview_parity(payload or {})

    # 3) Títulos já “Título, Aula N”
    try:
        p["titulos"] = _titles_for_export(p)
    except Exception:
        p["titulos"] = []

    # 4) BNCC consolidado quando textarea vazio (mesma lógica do DOCX)
    bncc_txt = p.get("conteudos_habilidades_bncc") or ""
    if not bncc_txt:
        linhas = p.get("linhas") or []
        linhas_txt = []
        for ln in linhas:
            try:
                codigo = (ln.get("codigo") or "").strip()
                habtx  = (ln.get("habilidade") or "").strip()
                if codigo or habtx:
                    linhas_txt.append(f"{codigo} {habtx}".strip())
            except Exception:
                continue
        bncc_txt = "\n".join(_unique_preserve_order(linhas_txt))
    p["conteudos_habilidades_bncc"] = bncc_txt

    return JSONResponse(p)


# [PATCH-EXPORT_PDF: HEADER] — Cabeçalho do /export/pdf com compat, paridade e helpers locais
@app.post("/export/pdf")
def export_pdf(payload: Dict[str, Any] = Body(...)):
    """
    Gera PDF padronizado do Plano de Aula.
    - Aceita payload no modelo novo (ExportPayload) e adapta para legado.
    - Garante paridade com a prévia (mesmas regras do DOCX).
    """
    # 1) Import guard do ReportLab
    try:
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.units import mm
    except Exception:
        raise HTTPException(status_code=501, detail="reportlab não instalado. Instale com: pip install reportlab")

    # 2) Compat: aceitar novo payload e adaptar (uma única vez)
    parsed = None
    if isinstance(payload, dict):
        try:
            parsed  = ExportPayload(**payload)
            adapted = _to_legacy_dict(parsed)
            payload = {**payload, **adapted}
            if getattr(parsed, "aulas", None):
                payload["aulas"] = parsed.aulas
        except Exception:
            parsed = None

    # 3) Paridade com a prévia (regras do DOCX)
    p = _ensure_preview_parity(payload or {})

    # 4) Helpers locais
    

    def _flat_lines(x):
        if x is None:
            return []
        if isinstance(x, (list, tuple, set)):
            out = []
            for it in x:
                s = str(it).strip()
                if s:
                    out.extend([t for t in s.split("\n") if t.strip()])
            return out
        return [t for t in str(x).split("\n") if t.strip()]

    def _esc(txt: str) -> str:
        if not isinstance(txt, str):
            txt = str(txt)
        txt = txt.replace("<b>", "§§B_OPEN§§").replace("</b>", "§§B_CLOSE§§")
        txt = _xml_escape(txt)
        return txt.replace("§§B_OPEN§§", "<b>").replace("§§B_CLOSE§§", "</b>")

    def _norm_no_accents(s: str) -> str:
        import unicodedata, re as _re
        s = unicodedata.normalize("NFD", str(s or ""))
        s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
        return _re.sub(r"\s+", " ", s).strip().casefold()

    COD_RE_INLINE = re.compile(r'[A-Z]{2}\d{2}[A-Z]{2}\d{2}')
    COD_RE_PREFIX = re.compile(r'^\s*\(?([A-Z]{2}\d{2}[A-Z]{2}\d{2})\)?[:\-\s]*', re.I)

    def _dedup_bncc_lines(lines_or_text) -> list[str]:
        lines = _flat_lines(lines_or_text)
        seen, out = set(), []
        for raw in lines:
            s = str(raw or "").strip()
            if not s:
                continue
            m = COD_RE_PREFIX.match(s) or COD_RE_INLINE.search(s)
            code = (m.group(1) if (m and m.groups()) else (m.group(0) if m else "")).strip() if m else ""
            desc = COD_RE_PREFIX.sub("", s).strip()
            key = (code.upper(), _norm_no_accents(desc))
            if key in seen:
                continue
            seen.add(key)
            out.append(f"({code}) {desc}" if code else desc)
        return out

    # 5) Derivações já com as mesmas regras do DOCX
    ident = p.get("identificacao") or "Plano de Aula"
    comps = ", ".join(_flat_lines(p.get("componente_curricular") or p.get("disciplina")))
    temas = ", ".join(_flat_lines(p.get("tema") or p.get("temas")))
    try:
        titulos_pdf = _titles_for_export(p)
    except Exception:
        titulos_pdf = []

    # 1) Habilidades: primeiro o textarea específico
    bncc_txt = p.get("habilidades_texto") or ""
    if not bncc_txt:
        bncc_txt = "\n".join(_flat_lines(p.get("habilidades_bncc")))

    # 2) Se vazio, cai no campo combinado (compat)
    if not bncc_txt:
        bncc_txt = p.get("conteudos_habilidades_bncc") or ""

    # 3) Se ainda vazio, fallback nas linhas cruas
    if not bncc_txt:
        linhas = p.get("linhas") or []
        linhas_txt = []
        for ln in linhas:
            try:
                codigo = (ln.get("codigo") or "").strip()
                habtx  = (ln.get("habilidade") or "").strip()
                if codigo or habtx:
                    linhas_txt.append((f"({codigo}) {habtx}".strip() if codigo else habtx))
            except Exception:
                continue
        bncc_txt = "\n".join(linhas_txt)

    bncc_list = _dedup_bncc_lines(bncc_txt)


    # 6) OA agrupados (defensivo)
    try:
        oa_grouped = _group_oa_by_hab_using_json(p) or {}
    except Exception:
        oa_grouped = {}

    obj_espec  = p.get("objetivos_especificos") or ""
    previos    = p.get("conhecimentos_previos") or ""
    metod      = p.get("metodologia_estrategias") or ""
    recs       = p.get("recursos_didaticos") or ""
    atividades = p.get("atividades_desempenho") or ""
    criterios  = p.get("criterios_avaliacao") or ""
    autoav     = p.get("avaliacao_autoavaliacao") or ""
    consid     = p.get("consideracoes_finais") or ""
    aulas_hab  = p.get("aulas_por_habilidade") or {}
    aulas_tp   = p.get("aulas_por_habilidade_tp") or {}
    tot_aulas  = p.get("total_aulas")
    tot_t      = p.get("total_aulas_teoricas")
    tot_p      = p.get("total_aulas_praticas")
    tec        = p.get("dados_tecnicos") or {}

    # 7) PDF / Fontes e estilos
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm, topMargin=18 * mm, bottomMargin=18 * mm
    )

    # Tenta fontes locais; se não houver APP_ROOT, cai para Helvetica
    base_font = "Helvetica"
    bold_font = None
    try:
        root = APP_ROOT if "APP_ROOT" in globals() else Path(".")
        arial = (root / "app" / "fonts" / "Arial.ttf")
        dejavu = (root / "app" / "fonts" / "DejaVuSans.ttf")
        if arial.exists():
            pdfmetrics.registerFont(TTFont("Arial", str(arial)))
            base_font = "Arial"
            arial_bold = (root / "app" / "fonts" / "Arial-Bold.ttf")
            if arial_bold.exists():
                pdfmetrics.registerFont(TTFont("Arial-Bold", str(arial_bold)))
                bold_font = "Arial-Bold"
        elif dejavu.exists():
            pdfmetrics.registerFont(TTFont("DejaVu", str(dejavu)))
            base_font = "DejaVu"
            dejavu_bold = (root / "app" / "fonts" / "DejaVuSans-Bold.ttf")
            if dejavu_bold.exists():
                pdfmetrics.registerFont(TTFont("DejaVu-Bold", str(dejavu_bold)))
                bold_font = "DejaVu-Bold"
    except Exception:
        base_font, bold_font = "Helvetica", None

    title_font = bold_font or base_font
    h1_font    = bold_font or base_font
    body_font  = base_font

    styles = getSampleStyleSheet()
    # Evita colisão, recriando nomes
    for nm in ("TitleX", "H1X", "P"):
        if nm in styles.byName:
            del styles.byName[nm]
    styles.add(ParagraphStyle(
        name="TitleX", parent=styles["Title"],
        fontName=title_font, fontSize=16, leading=20,
        alignment=TA_CENTER, spaceAfter=8
    ))
    styles.add(ParagraphStyle(
        name="H1X", parent=styles["Heading2"],
        fontName=h1_font, fontSize=12, leading=16,
        spaceBefore=8, spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name="P", parent=styles["BodyText"],
        fontName=body_font, fontSize=12, leading=15, alignment=TA_LEFT
    ))

    story: List[Any] = []

    def add_h1(txt: str):
        story.append(Paragraph(_esc(f"<b>{txt}</b>"), styles["H1X"]))
        story.append(Spacer(1, 6))

    def add_list(items):
        _items = _flat_lines(items) or ["—"]
        bullets = [ListItem(Paragraph(_esc(str(it)), styles["P"])) for it in _items]
        story.append(ListFlowable(bullets, bulletType="bullet", start="•", leftPadding=12))
        story.append(Spacer(1, 6))

    # 8) Montagem
    story.append(Paragraph(_esc("Plano de Aula"), styles["TitleX"]))
    story.append(Spacer(1, 8))

    # Identificação
    _pl = p.get("plano") or {}
    escola    = (p.get("identificacao") or "").strip()
    professor = _pl.get("professor")  or p.get("professor")  or ""
    turma     = _pl.get("turma")      or p.get("turma")      or ""
    data_     = _pl.get("data")       or p.get("data")       or ""
    ano       = _pl.get("ano")        or p.get("ano")        or ""
    bimestre  = _pl.get("bimestre")   or p.get("bimestre")   or ""
    etapa_hdr = _pl.get("etapa")      or p.get("etapa")      or ""

    id_lines = []
    if escola:    id_lines.append(f"<b>Escola:</b> {escola}")
    if professor: id_lines.append(f"<b>Professor(a):</b> {professor}")
    if turma:     id_lines.append(f"<b>Turma:</b> {turma}")
    if ano:       id_lines.append(f"<b>Ano letivo:</b> {ano}")
    if bimestre:  id_lines.append(f"<b>Bimestre:</b> {bimestre}")
    if etapa_hdr: id_lines.append(f"<b>Etapa:</b> {etapa_hdr}")
    for line in id_lines:
        story.append(Paragraph(_esc(line), styles["P"]))
    if id_lines:
        story.append(Spacer(1, 8))

    meta_lines = []
    if comps: meta_lines.append(f"<b>Componente(s):</b> {comps}")
    if temas: meta_lines.append(f"<b>Tema(s):</b> {temas}")
    for m in meta_lines:
        story.append(Paragraph(_esc(m), styles["P"]))
    if meta_lines:
        story.append(Spacer(1, 6))

    # Habilidades BNCC
    if bncc_list:
        add_h1("Habilidades (BNCC)")
        for line in bncc_list:
            story.append(Paragraph(_esc(line), styles["P"]))
        story.append(Spacer(1, 6))

    # Títulos
    if titulos_pdf:
        add_h1("Título(s) da Aula")
        add_list(titulos_pdf)

    # Objetivos Específicos
    if _flat_lines(obj_espec):
        add_h1("Objetivos Específicos")
        add_list(obj_espec)

    # Objetivos de Aprendizagem
    if oa_grouped:
        add_h1("Objetivos de Aprendizagem")
        for code, linhas in oa_grouped.items():
            line = f"<b>{code}:</b> " + " — ".join(_flat_lines(linhas))
            story.append(Paragraph(_esc(line), styles["P"]))
            story.append(Spacer(1, 6))

    # Conteúdos
    conts = _flat_lines(p.get("conteudos"))
    if conts:
        add_h1("Conteúdos")
        add_list(conts)

    # Demais seções
    if _flat_lines(previos):
        add_h1("Conhecimentos Prévios");      add_list(previos)
    if _flat_lines(metod):
        add_h1("Metodologias e Estratégias"); add_list(metod)
    if _flat_lines(recs):
        add_h1("Recursos Didáticos");         add_list(recs)
    if _flat_lines(atividades):
        add_h1("Atividades/Desempenho");      add_list(atividades)
    if _flat_lines(criterios):
        add_h1("Critérios de Avaliação");     add_list(criterios)
    if _flat_lines(autoav):
        add_h1("Avaliação/Autoavaliação");    add_list(autoav)
    if _flat_lines(consid):
        add_h1("Considerações Finais");       add_list(consid)

    # Aulas por Habilidade
    if aulas_hab:
        add_h1("Aulas por Habilidade")
        lines = [f"{k}: {v} aula(s)" for k, v in aulas_hab.items()]
        add_list(lines)
        if tot_aulas is not None:
            add_list([f"Total de aulas: {tot_aulas}"])

    # Técnico (T/P)
    if aulas_tp or (tot_t is not None) or (tot_p is not None):
        add_h1("Aulas por Habilidade (Técnico)")
        if aulas_tp:
            lines = [f"{k}: {vp.get('t',0)} T / {vp.get('p',0)} P" for k, vp in aulas_tp.items()]
            add_list(lines)
        if (tot_t is not None) or (tot_p is not None):
            soma = (tot_t or 0) + (tot_p or 0)
            add_list([f"Total: {tot_t or 0} T / {tot_p or 0} P (Soma: {soma})"])

    # Campos Técnicos
    if any(tec.values()) if isinstance(tec, dict) else False:
        add_h1("Campos Técnicos")
        lines = []
        if tec.get("ambiente"):      lines.append("Ambiente/Lab: " + ", ".join(_flat_lines(tec["ambiente"])))
        if tec.get("epis"):          lines.append("Equip/EPIs: "  + ", ".join(_flat_lines(tec["epis"])))
        if tec.get("seguranca"):     lines.append("Segurança: "   + ", ".join(_flat_lines(tec["seguranca"])))
        if tec.get("requisitos"):    lines.append("Requisitos: "  + ", ".join(_flat_lines(tec["requisitos"])))
        if tec.get("procedimentos"): lines.append("Procedimentos: "+ ", ".join(_flat_lines(tec["procedimentos"])))
        add_list(lines or ["—"])

    # 9) Constrói e retorna
    doc.build(story)
    filename = (_slug(ident) or "plano") + ".pdf"
    return Response(
        content=buf.getvalue(),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        media_type="application/pdf",
    )



@app.post("/export/xlsx")
def export_xlsx(payload: Dict[str, Any] = Body(...)):
    """Gera planilha Excel com abas por seção do plano (paridade com prévia/DOCX/PDF)."""

    # Aceita novo payload sem quebrar o legado
    if isinstance(payload, dict):
        try:
            parsed  = ExportPayload(**payload)
            adapted = _to_legacy_dict(parsed)
            payload = {**payload, **adapted}
        except Exception:
            pass

    payload = _ensure_preview_parity(payload or {})

    try:
        import pandas as pd
    except Exception:
        raise HTTPException(
            status_code=501,
            detail="Dependências ausentes: instale com `pip install pandas xlsxwriter`."
        )


    ident = payload.get("identificacao") or "Plano de Aula"
    meta = {
        "Identificação": [ident],
        "Componente Curricular": [", ".join(_flatten_lines(payload.get("componente_curricular")))],
        "Tema": [", ".join(_flatten_lines(payload.get("tema")))],
        "Títulos da Aula": [", ".join(_flatten_lines(_titles_for_export(payload)))],
    }

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as w:
        # 00_Meta
        pd.DataFrame(meta).to_excel(w, index=False, sheet_name="00_Meta")

        # Aux: gravar lista simples (uma coluna)
        def tab(nome: str, items):
            df = pd.DataFrame({nome: _flatten_lines(items)})
            df.to_excel(w, index=False, sheet_name=nome[:31])

        tab("01_HabilidadesBNCC", payload.get("conteudos_habilidades_bncc"))
        tab("02_ObjetivosAprend", payload.get("objetivos_aprendizagem"))
        tab("03_ObjetivosEspec",  payload.get("objetivos_especificos"))
        tab("04_ConhecPrevios",   payload.get("conhecimentos_previos"))
        tab("05_Metodologia",     payload.get("metodologia_estrategias"))
        tab("06_Recursos",        payload.get("recursos_didaticos"))
        tab("07_Atividades",      payload.get("atividades_desempenho"))
        tab("08_CriteriosAvali",  payload.get("criterios_avaliacao"))
        tab("09_Autoavaliacao",   payload.get("avaliacao_autoavaliacao"))
        tab("10_Consideracoes",   payload.get("consideracoes_finais"))

        # 11_AulasPorHab (geral)
        aulas = payload.get("aulas_por_habilidade") or {}
        if aulas:
            df = pd.DataFrame({"Habilidade": list(aulas.keys()), "Aulas": list(aulas.values())})
            df.to_excel(w, index=False, sheet_name="11_AulasPorHab")

        # 12_AulasPorHab_TP (técnico)
        aulas_tp = payload.get("aulas_por_habilidade_tp") or {}
        if aulas_tp:
            df = pd.DataFrame({
                "Habilidade": list(aulas_tp.keys()),
                "Teoricas":   [aulas_tp[k].get("t",0) for k in aulas_tp],
                "Praticas":   [aulas_tp[k].get("p",0) for k in aulas_tp]
            })
            df.to_excel(w, index=False, sheet_name="12_AulasPorHab_TP")

        # 13_Tecnico (campos extras)
        tec = payload.get("dados_tecnicos") or {}
        if any(tec.values()):
            df = pd.DataFrame({
                "Ambiente/Lab": _flatten_lines(tec.get("ambiente")),
                "Equip/EPIs":   _flatten_lines(tec.get("epis")),
                "Segurança":    _flatten_lines(tec.get("seguranca")),
                "Requisitos":   _flatten_lines(tec.get("requisitos")),
                "Procedimentos":_flatten_lines(tec.get("procedimentos")),
            })
            if df.empty:
                df = pd.DataFrame({"Info": ["(sem dados)"]})
            df.to_excel(w, index=False, sheet_name="13_Tecnico")

    return Response(
        content=buf.getvalue(),
        headers={"Content-Disposition": f'attachment; filename="{_slug(ident) or "plano"}.xlsx"'},
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


PID_RE = re.compile(r"^[a-f0-9]{16}$")

def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds").replace("+00:00", "Z")

def _safe_slug(s: str, fallback: str = "plano") -> str:
    try:
        return _slug(s) or fallback
    except Exception:
        return fallback

def _assert_valid_pid(pid: str):
    if not PID_RE.match(pid or ""):
        raise HTTPException(status_code=400, detail="pid inválido")

def _user_dir_for(user_id: str) -> Path:
    uid = _slug_user(user_id or "anon")
    d = SAVED_DIR / uid
    d.mkdir(parents=True, exist_ok=True)
    return d

def _atomic_write(path: Path, content: str):
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(content, encoding="utf-8")
    tmp.replace(path)

# ---------- CREATE ----------
@app.post("/api/plans")
def create_plan(body: Dict[str, Any] = Body(...)):
    # NEW: valida payload ANTES de usar
    payload = body.get("payload")
    if not isinstance(payload, dict):
        raise HTTPException(status_code=422, detail="payload inválido")

    user_id = str(body.get("user_id") or "anon")
    title_in = str(body.get("title") or "").strip()
    title = title_in or str(payload.get("identificacao") or "Plano").strip()  # NEW: agora seguro
    title = _safe_slug(title, "Plano")

    udir = _user_dir_for(user_id)
    now = _utc_now_iso()

    raw_id = f"{_slug_user(user_id)}:{title}:{now}"
    pid = hashlib.sha1(raw_id.encode("utf-8")).hexdigest()[:16]

    doc = {
        "id": pid,
        "user_id": _slug_user(user_id),
        "title": title,
        "created_at": now,
        "payload": payload
    }

    p = udir / f"{pid}.json"
    _atomic_write(p, json.dumps(doc, ensure_ascii=False, indent=2))
    return {"id": pid, "ok": True, "created_at": now}

# ---------- LIST ----------
@app.get("/api/plans")
def list_plans(
    user_id: str = Query(...),
    limit: int = Query(100, ge=1, le=500),
    cursor: Optional[str] = Query(None, description="ISO-8601 created_at cursor; retorna < cursor"),
):
    udir = _user_dir_for(user_id)
    if not udir.exists():
        return {"items": [], "next_cursor": None}

    items = []
    # Ordena por mtime decrescente (mais novo primeiro)
    files = sorted(udir.glob("*.json"), key=lambda x: x.stat().st_mtime, reverse=True)

    for f in files:
        try:
            j = json.loads(f.read_text(encoding="utf-8"))
            created = j.get("created_at") or ""
            if cursor and created >= cursor:
                # queremos apenas ESTRITAMENTE menores que o cursor
                continue
            items.append({
                "id": j.get("id"),
                "title": j.get("title"),
                "created_at": created,
                "size": f.stat().st_size
            })
            if len(items) >= limit:
                break
        except Exception:
            continue

    next_cursor = items[-1]["created_at"] if len(items) == limit else None
    return {"items": items, "next_cursor": next_cursor}

# ---------- GET ----------
@app.get("/api/plans/{pid}")
def get_plan(pid: str, user_id: str = Query(...)):  # NEW: escopo por usuário
    _assert_valid_pid(pid)
    udir = _user_dir_for(user_id)
    p = udir / f"{pid}.json"
    if not p.exists():
        raise HTTPException(status_code=404, detail="Plano não encontrado")
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(status_code=422, detail="JSON do plano está corrompido.")

# ---------- DOWNLOAD ----------
@app.get("/api/plans/{pid}/download")
def download_plan(pid: str, user_id: str = Query(...)):  # NEW: escopo por usuário
    _assert_valid_pid(pid)
    udir = _user_dir_for(user_id)
    p = udir / f"{pid}.json"
    if not p.exists():
        raise HTTPException(status_code=404, detail="Plano não encontrado")
    return FileResponse(
        path=p,
        media_type="application/json",
        filename=f"plano_{pid}.json"
    )

# ---------- DELETE ----------
@app.delete("/api/plans/{pid}")
def delete_plan(pid: str, user_id: str = Query(...)):  # NEW: escopo por usuário
    _assert_valid_pid(pid)
    udir = _user_dir_for(user_id)
    p = udir / f"{pid}.json"
    if not p.exists():
        raise HTTPException(status_code=404, detail="Plano não encontrado")
    p.unlink(missing_ok=False)
    return {"ok": True}

# ---------- UPDATE ----------
@app.put("/api/plans/{pid}")
def update_plan(pid: str, body: Dict[str, Any] = Body(...), user_id: str = Query(...)):  # NEW: escopo por usuário
    """
    Atualiza um plano existente (payload/title). Mantém o mesmo id.
    """
    _assert_valid_pid(pid)
    udir = _user_dir_for(user_id)
    p = udir / f"{pid}.json"
    if not p.exists():
        raise HTTPException(status_code=404, detail="Plano não encontrado para atualização")

    try:
        doc = json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(status_code=422, detail="JSON do plano está corrompido.")

    updated = False

    # title
    if "title" in body:
        t = str(body.get("title") or "").strip()
        if t:
            doc["title"] = _safe_slug(t, doc.get("title") or "Plano")
            updated = True

    # payload
    if "payload" in body:
        pl = body.get("payload")
        if isinstance(pl, dict):
            doc["payload"] = pl
            updated = True
        elif pl is not None:
            raise HTTPException(status_code=422, detail="payload deve ser um objeto JSON")

    if not updated:
        return {"ok": True, "id": pid, "message": "Nada a atualizar."}

    doc["updated_at"] = _utc_now_iso()
    _atomic_write(p, json.dumps(doc, ensure_ascii=False, indent=2))
    return {"ok": True, "id": pid, "updated_at": doc["updated_at"]}